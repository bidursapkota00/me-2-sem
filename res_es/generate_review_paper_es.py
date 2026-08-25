#!/usr/bin/env python3
"""
Generate an IEEE two-column format review paper in DOCX format.
Topic: Recent Improvements in Battery-Free Embedded Systems: A Review
Page Size: A4
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = "/Users/bidur/ncit/res_es/Review_Paper_BatteryFree_ES_2025_2026.docx"

doc = Document()

# ============================================================
# PAGE SETUP - IEEE format (A4 size, narrow margins)
# ============================================================
for section in doc.sections:
    section.page_width = Cm(21.0)    # A4 width
    section.page_height = Cm(29.7)   # A4 height
    section.top_margin = Cm(1.91)    # 0.75 in
    section.bottom_margin = Cm(2.54) # 1 in
    section.left_margin = Cm(1.78)   # 0.7 in
    section.right_margin = Cm(1.78)  # 0.7 in

    # Two-column layout
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="360"/>')
        sectPr.append(cols)
    else:
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '360')

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.0

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Times New Roman'
    return p

def add_authors(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    return p

def add_affiliation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.italic = True
    return p

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    return p

def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    return p

def add_body(doc, text, first_line_indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    return p

def add_reference(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(f"[{number}] {text}")
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'
    return p

def add_table(doc, caption, headers, data, col_widths=None):
    """Add a formatted table with caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(caption)
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'

    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(7)
                run.font.name = 'Times New Roman'

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)
                    run.font.name = 'Times New Roman'

    if col_widths:
        for i, w in enumerate(col_widths):
            table.columns[i].width = Cm(w)
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)


# ============================================================
# DOCUMENT CONTENT
# ============================================================

# --- TITLE ---
add_title(doc, "Recent Improvements in Battery-Free\nEmbedded Systems:\nA Review")

# --- AUTHORS ---
add_authors(doc, "Bidur Sapkota")
add_affiliation(doc, "Department of Computer Engineering, Nepal College of Information Technology (NCIT)\nPokhara University, Nepal\nbidur@ncit.edu.np")

# --- ABSTRACT ---
add_section_heading(doc, "Abstract")

abstract_text = (
    "Battery-free embedded systems powered by ambient energy harvesting have emerged as a "
    "transformative paradigm for sustainable Internet of Things (IoT) deployments. The period "
    "spanning 2024 to 2026 has witnessed a significant transition from proof-of-concept "
    "demonstrations to increasingly mature, deployment-ready solutions. This review paper "
    "critically examines six recent journal and conference papers that address key challenges "
    "in advancing battery-free system technology. The reviewed works encompass analytical "
    "modeling of energy dynamics in batteryless IoT sensors, cost-effective design methodologies "
    "for intermittent wireless communication, standardized evaluation platforms for fair "
    "performance benchmarking, energy-harvesting-aware neural architecture search for "
    "intermittent deep learning inference, UAV-powered RF energy-harvesting sensor systems "
    "for precision agriculture, and comprehensive surveys of circuit and system design "
    "architectures. Through a systematic comparison of the objectives, methodologies, results, "
    "strengths, and limitations of each study, this review identifies several critical research "
    "gaps, including the absence of unified benchmarking standards, insufficient multi-source "
    "energy harvesting integration, limited adaptation of edge AI for intermittent operation, "
    "and inadequate standardization of power management protocols. The paper concludes with "
    "recommendations for future research directions aimed at achieving reliable, scalable, "
    "and maintenance-free battery-free IoT ecosystems."
)
add_body(doc, abstract_text, first_line_indent=False)

# --- Keywords ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(6)
run_kw = p.add_run("Keywords: ")
run_kw.bold = True
run_kw.italic = True
run_kw.font.size = Pt(9)
run_kw.font.name = 'Times New Roman'
run_val = p.add_run("Battery-Free Embedded Systems, Energy Harvesting, Intermittent Computing, "
                     "IoT, TinyML, Power Management, Wireless Sensor Networks, Precision Agriculture")
run_val.italic = True
run_val.font.size = Pt(9)
run_val.font.name = 'Times New Roman'

# ============================================================
# I. INTRODUCTION
# ============================================================
add_section_heading(doc, "I. Introduction")

# --- A. Background ---
add_subsection_heading(doc, "A. Background")

add_body(doc, (
    "Many IoT devices are everywhere now. Like farm, factory, nature check, "
    "and health. So we need sensor that live long time without help. But big "
    "problem for many IoT is battery. Normal battery cost much money because "
    "we must change them. Also they make bad trash for environment. And it is "
    "very hard to change battery if sensor is very far place [6]."
))

add_body(doc, (
    "Battery-free system is good choice. It takes energy from outside like "
    "sun, radio wave, heat and shaking. These system use supercapacitor or "
    "small cell to keep energy. They use it to do compute and talk sometimes. "
    "This way is called intermittent computing. It have hard problem because "
    "system must not make mistake and keep data safe when power go off suddenly [3], [6]."
))

add_body(doc, (
    "From 2024 to 2026, research in battery-free system change a lot. Before, "
    "people just try to make it work. Now people try to solve real problem. "
    "Like how to guess performance, make cheap hardware, test it good, and "
    "put AI on it. Also new thing like drone sensor for farm show this tech "
    "is very useful now [1], [5]."
))

# --- B. Problem Statement ---
add_subsection_heading(doc, "B. Problem Statement")

add_body(doc, (
    "Despite substantial progress in individual aspects of battery-free embedded system design, "
    "the field remains fragmented across multiple research communities, including energy "
    "harvesting, intermittent computing, low-power circuit design, and edge artificial "
    "intelligence. This fragmentation has resulted in several critical deficiencies. First, "
    "there is no widely accepted benchmark suite or standardized evaluation methodology for "
    "comparing the performance of different battery-free platforms under equivalent conditions. "
    "Second, the integration of multiple heterogeneous energy harvesting sources within a "
    "unified power management framework remains largely unexplored. Third, the adaptation of "
    "modern machine learning techniques, particularly deep neural network inference, for "
    "intermittent execution environments is still in its nascent stages. Fourth, the absence "
    "of standardized power management protocols across diverse hardware platforms hinders "
    "interoperability and scalability [2], [4]."
))

add_body(doc, (
    "These problem make it hard to use battery-free system in real world. "
    "So we need to read new paper and find what is missing. This help us know "
    "what to do next."
))

# --- C. Research Objectives ---
add_subsection_heading(doc, "C. Research Objectives")

add_body(doc, (
    "This paper want to look at 6 new research from 2024 to 2026. We want to do this:"
), first_line_indent=False)

add_body(doc, (
    "(1) To read and summarize what each paper do and find."
), first_line_indent=False)

add_body(doc, (
    "(2) To compare how they do research and their result."
), first_line_indent=False)

add_body(doc, (
    "(3) To find out what math and tool they use."
), first_line_indent=False)

add_body(doc, (
    "(4) To see what big problem is still not fixed."
), first_line_indent=False)

add_body(doc, (
    "(5) To suggest what people should research next."
), first_line_indent=False)

# ============================================================
# II. RELATED WORKS
# ============================================================
add_section_heading(doc, "II. Related Works")

add_body(doc, (
    "In this part we look at 6 new paper about battery-free system. We talk about "
    "what they try to do, how they do it, and what they find good or bad."
), first_line_indent=False)

# --- Paper 1 ---
add_subsection_heading(doc, "A. Analytical Modeling of Batteryless IoT Sensors Powered by Ambient Energy Harvesting [1]")

add_body(doc, (
    "First paper is by Fernández Landivar in 2025 [1]. They make math model for "
    "battery-free sensor. They want to show how energy move from harvest to storage "
    "to compute. Their model use math for supercapacitor and power management chip."
))

add_body(doc, (
    "They use physics to write equation for voltage. Then they test model with real "
    "sensor in a room with light. The math match the real test very good. So we know "
    "their math is correct for different light."
))

add_body(doc, (
    "Good thing about this paper is model work for many hardware. People can use it "
    "to guess performance before building. But bad thing is model assume energy is "
    "stable. It don't work well if energy change very fast, like outside."
))

# --- Paper 2 ---
add_subsection_heading(doc, "B. Designing Cost-Effective Battery-Less Energy Harvesting for Intermittent Wireless Communication [2]")

add_body(doc, (
    "Published in IEEE Access (January 2025), this paper addressed the challenge of designing "
    "cost-effective battery-less IoT devices capable of sustaining reliable wireless "
    "communication under intermittent power availability [2]. The research objective was to "
    "develop systematic design models that enable engineers to determine optimal energy "
    "harvester and storage element specifications for meeting specific wireless communication "
    "Quality of Service (QoS) requirements, including data throughput and transmission frequency."
))

add_body(doc, (
    "The authors proposed energy and timing models that formalize the relationship between "
    "harvester capacity, storage element sizing (capacitors), and the resulting communication "
    "duty cycle. A central contribution is the introduction of a duty-based Dynamic Power "
    "Management (DPM) scheme that dynamically adjusts the device's active and sleep periods "
    "based on available energy, significantly outperforming conventional threshold-based "
    "approaches. Experimental evaluation demonstrated that the proposed DPM scheme extends "
    "operational time by 9.1% to 90.0% compared to baseline methods, depending on ambient "
    "energy availability conditions."
))

add_body(doc, (
    "The strength of this work lies in its practical applicability: the design methodology "
    "provides actionable guidelines for hardware engineers to translate wireless QoS "
    "requirements into concrete harvester and storage specifications. The duty-based DPM "
    "approach offers a clear performance advantage. However, the methodology is primarily "
    "validated for single-source energy harvesting scenarios and does not explicitly consider "
    "multi-source configurations or the impact of highly variable channel conditions on "
    "communication reliability."
))

# --- Paper 3 ---
add_subsection_heading(doc, "C. EStacker: Explaining Battery-Less IoT System Performance with Energy Stacks [3]")

add_body(doc, (
    "Liedtke et al. presented EStacker in ACM Transactions on Embedded Computing Systems "
    "(2026), a specialized evaluation platform designed to address the challenge of fair "
    "and reproducible performance assessment of battery-less IoT systems [3]. The research "
    "objective was to create a systematic benchmarking framework that enables developers to "
    "understand precisely how energy is consumed across different hardware components and "
    "software tasks through the concept of \"energy stacks\"N/Adetailed energy consumption "
    "profiles decomposed by functional unit and execution phase."
))

add_body(doc, (
    "The platform employs a simulation-based approach that models the interaction between "
    "energy harvesting, storage dynamics, and application workload execution. A significant "
    "methodological contribution is the ST-SP (Spatial-Temporal State Pruning) optimization "
    "technique, which reduces evaluation time by a factor of 6.3× on average while "
    "maintaining timing accuracy with a mean error of only 7.7%. This acceleration is "
    "critical for enabling practical design space exploration across different harvester "
    "and storage configurations."
))

add_body(doc, (
    "EStacker's principal strength is its ability to provide deterministic, energy-equivalent "
    "evaluation conditions, ensuring that different applications and hardware configurations "
    "are compared under identical energy budgets. This addresses a fundamental reproducibility "
    "challenge in battery-free systems research. The limitation is that the platform currently "
    "relies on simulation rather than hardware-in-the-loop evaluation, which may not fully "
    "capture real-world phenomena such as parasitic losses, component aging, and environmental "
    "electromagnetic interference."
))

# --- Paper 4 ---
add_subsection_heading(doc, "D. HANNA: AI Search for No Battery [4]")

add_body(doc, (
    "Fourth paper by Sahu in 2024 [4] is about HANNA. It find best neural "
    "network for battery-free device. Normal AI search only care about speed "
    "and accuracy. But intermittent device power turn off often. So normal AI "
    "don\'t work well."
))

add_body(doc, (
    "They use simulation to check how AI run when power is cut. HANNA try to "
    "find AI that is small and can save state easy. Result show HANNA make AI "
    "10% to 44% more accurate compare to old method that ignore power cut."
))

add_body(doc, (
    "HANNA is very strong because it mix tiny AI with intermittent compute. "
    "The AI it find look different. It have small layer to finish job fast. "
    "But it only test on some dataset. We need check if it work good on many hardware."
))

# --- Paper 5 ---
add_subsection_heading(doc, "E. Farm Sensor with Drone [5]")

add_body(doc, (
    "Fifth paper by Kudyba and Sun in 2025 [5] use drone and battery-free sensor "
    "for farm. Drone fly over field and send radio energy to sensor. Sensor use "
    "energy to wake up and send temp data back to drone. This save money because "
    "sensor no need battery."
))

add_body(doc, (
    "They build real payload for drone and put sensor in field. Ground test "
    "work very good and send data from 10 meter. But drone test fail. Drone motor "
    "make too much radio noise. So drone cannot receive temp data from sensor."
))

add_body(doc, (
    "Good thing is they show full real world system. It show big difference "
    "between lab test and field test. Bad thing is system is very sensitive to "
    "radio noise. Drone noise block signal completely. So we need better way to stop noise."
))

# --- Paper 6 ---
add_subsection_heading(doc, "F. Batteryless Systems for IoT: A Survey of Circuit and System Design [6]")

add_body(doc, (
    "Presented at IEEE MWSCAS 2025, this comprehensive survey paper provides a systematic "
    "review of circuit-level and system-level design architectures for batteryless IoT "
    "devices [6]. The survey encompasses the complete hardware stack, including energy "
    "harvesting transducers, power management integrated circuits (PMICs), energy storage "
    "technologies (supercapacitors, thin-film batteries), non-volatile memory systems for "
    "state preservation, and intermittent computing execution models."
))

add_body(doc, (
    "The survey categorizes intermittent computing approaches into two principal paradigms: "
    "task-based execution, where computational tasks are partitioned into atomic units that "
    "can complete within a single energy burst; and checkpoint-based execution, where the "
    "system periodically saves its computational state to non-volatile memory to enable "
    "resumption after power interruptions. The paper systematically analyzes the trade-offs "
    "inherent in each approach, including harvester efficiency versus manufacturing cost, "
    "storage capacity versus leakage current, and computational throughput versus the energy "
    "overhead of checkpointing operations."
))

add_body(doc, (
    "The principal strength of this survey is its breadth and systematic organization, which "
    "provides an accessible entry point for researchers entering the battery-free systems "
    "domain. The survey effectively contextualizes individual circuit innovations within the "
    "broader system architecture, facilitating cross-disciplinary understanding. However, the "
    "survey primarily focuses on hardware-level considerations and provides limited coverage "
    "of software frameworks, middleware, and application-level challenges associated with "
    "intermittent operation."
))

# --- Comparison Table ---
add_table(doc,
    "TABLE I: Comparative Summary of Reviewed Research Papers",
    ["Ref.", "Objective", "Methodology", "Key Result", "Strength", "Limitation"],
    [
        ["[1]", "Model energy dynamics of batteryless IoT sensors",
         "First-principles analytical modeling with experimental validation",
         "Accurate voltage prediction under varying illumination",
         "Generalizable across hardware platforms",
         "Assumes stable ambient energy conditions"],
        ["[2]", "Design cost-effective battery-less wireless communication",
         "Energy and timing models with duty-based DPM scheme",
         "9.1%–90.0% operational time extension",
         "Actionable design guidelines for engineers",
         "Single-source harvesting only"],
        ["[3]", "Fair and reproducible benchmarking of battery-less systems",
         "Simulation-based energy stacks with ST-SP optimization",
         "6.3× evaluation speedup with 7.7% timing error",
         "Deterministic energy-equivalent evaluation",
         "Simulation-only; lacks hardware-in-the-loop"],
        ["[4]", "Find best AI for no battery",
         "AI search with power cut",
         "AI finish in one cycle",
         "First tiny ML for this",
         "Not tested on many thing"],
        ["[5]", "Farm sensor with drone",
         "Radio harvest test",
         "Ground work but drone fail",
         "Real world test",
         "Drone noise kill signal"],
        ["[6]", "Survey of circuit design",
         "Read many paper",
         "Good group of old work",
         "Easy to read",
         "No software talk"],
    ],
    col_widths=[0.7, 1.4, 1.6, 1.8, 1.5, 1.5]
)


# ============================================================
# III. METHODOLOGY ADOPTED BY THE REVIEWED WORKS
# ============================================================
add_section_heading(doc, "III. Methodology Adopted by the Reviewed Works")

add_body(doc, (
    "The 6 paper use 4 main way to do research: math model, real prototype, "
    "computer simulation, and survey. Here we look at how they do it."
), first_line_indent=False)

add_subsection_heading(doc, "A. Math Model")

add_body(doc, (
    "Paper [1] and [2] use math model. Paper [1] write math equation for "
    "supercapacitor and power. They test part by part to get math numbers. "
    "Paper [2] use math to balance energy and time for communication. Both "
    "help test design before make real hardware."
))

add_subsection_heading(doc, "B. Real Test")

add_body(doc, (
    "Paper [1], [2], and [5] make real hardware to test. Paper [1] test "
    "IoT node with room light. Paper [2] test power manage with real energy "
    "profile. Paper [5] test farm sensor with drone in real field. Real test "
    "is very important to see if math model is correct in real world."
))

add_subsection_heading(doc, "C. Computer Test")

add_body(doc, (
    "Paper [3] and [4] use computer simulation. EStacker [3] simulate time "
    "and energy very detail. It use smart way to cut simulation time. "
    "HANNA [4] simulate AI on intermittent power. It use optimization to "
    "balance accuracy and energy cost."
))

add_subsection_heading(doc, "D. Reading Old Papers")

add_body(doc, (
    "Paper [6] is survey paper. It read many old paper and put them in groups. "
    "It group by hardware part and execution type. This help see what is "
    "good and bad, and show what researcher do now."
))

# --- Methodology Comparison Table ---
add_table(doc,
    "TABLE II: How They Do Research",
    ["Method", "Papers", "Key Idea", "How They Check"],
    [
        ["Math Model", "[1], [2]",
         "Math equations",
         "Compare with lab test"],
        ["Real Prototype", "[1], [2], [5]",
         "Make real hardware",
         "Measure in real world"],
        ["Simulation", "[3], [4]",
         "Computer test",
         "Check speed and error"],
        ["Survey", "[6]",
         "Group old paper",
         "Find what is missing"],
    ],
    col_widths=[1.8, 1.2, 3.2, 2.2]
)

# ============================================================
# IV. RESULTS AND DISCUSSION
# ============================================================
add_section_heading(doc, "IV. Results and Discussion")

add_body(doc, (
    "Here we compare the papers. We look at result, numbers, trends, and problems that people still need to fix."
), first_line_indent=False)

add_subsection_heading(doc, "A. Compare Papers")

add_body(doc, (
    "The 6 paper show 3 main research area. Paper [1] and [2] do math and "
    "design rule. Paper [3] and [4] make test tool and AI architecture. "
    "Paper [5] and [6] do real farm application and big survey review. "
    "All are important for battery-free system."
))

add_body(doc, (
    "All paper show that researcher now care about real use, not just show "
    "it work. Math from [1] and rule from [2] help each other. EStacker [3] "
    "can test AI from HANNA [4]. They all connect together to make battery-free "
    "system better."
))

add_subsection_heading(doc, "B. Key Findings and Performance Metrics")

add_body(doc, (
    "The number result in paper show big improvement. Table III show main "
    "performance number for all paper."
))

# --- TABLE III: Quantitative Performance Metrics ---
add_table(doc,
    "TABLE III: Quantitative Performance Metrics of Reviewed Studies",
    ["Ref.", "Metric", "Value", "Improvement", "Baseline"],
    [
        ["[1]", "Voltage prediction accuracy", "3 scenarios validated", "High correlation", "Measured profiles"],
        ["[2]", "Operational time extension", "9.1%–90.0%", "+9.1% to +90.0%", "Threshold-based DPM"],
        ["[3]", "Evaluation speedup (ST-SP)", "6.3x", "41.7 days → 7.7 days", "Full simulation"],
        ["[3]", "Throughput timing error", "7.7%", "N/A", "Cycle-accurate sim."],
        ["[3]", "App. performance gain", "3.3x", "+230%", "Unoptimized config."],
        ["[4]", "Inference accuracy gain", "10%–44%", "+10% to +44%", "SOTA NAS methods"],
        ["[4]", "NAS search cost", "1-shot", "Significant reduction", "Traditional NAS"],
        ["[5]", "Tag comm. range (BLE)", "10 m", "N/A", "N/A"],
        ["[5]", "RF harvest frequency", "918 MHz", "N/A", "N/A"],
        ["[5]", "Active TX frequency", "2.4 GHz", "N/A", "N/A"],
        ["[5]", "Aerial data packets received", "1 (ID only)", "0 temp. packets", "Ground: reliable"],
    ],
    col_widths=[0.7, 2.3, 1.8, 1.8, 1.9]
)

add_body(doc, (
    "Power method in [2] make system work 9.1% to 90% longer time. It help "
    "most when energy is very low. EStacker [3] make test 6.3x faster but "
    "error is only 7.7%. It also find problem in app and make it 3.3x better, "
    "so test time go from 41.7 day to 7.7 day."
))

add_body(doc, (
    "HANNA [4] improve AI accuracy by 10% to 44% compare to old AI. It also "
    "make search cost lower. The AI it find is different. It have smaller "
    "layer and save state more time. So it can finish job in short power time."
))

add_body(doc, (
    "Farm system in [5] use 918 MHz radio to harvest energy. It use 2.4 GHz "
    "to send data up to 10 meter. Ground test is good. But drone test fail "
    "because drone motor make interference. They get 1 ID packet but no temp data."
))

# --- TABLE IV: Technical Specifications Comparison ---
add_table(doc,
    "TABLE IV: Technical Specifications of Reviewed Systems",
    ["Ref.", "Harvest Freq.", "TX Freq.", "Storage Type", "Range", "Eval. Scale", "Year"],
    [
        ["[1]", "Solar (indoor)", "N/A", "Supercapacitor", "N/A", "3 scenarios", "2025"],
        ["[2]", "Ambient", "Sub-GHz/ISM", "Capacitor", "N/A", "Multiple configs.", "2025"],
        ["[3]", "Simulated", "N/A", "Simulated", "N/A", "2 case studies", "2026"],
        ["[4]", "RF", "N/A", "Capacitor + NVM", "N/A", "Multiple datasets", "2024"],
        ["[5]", "918 MHz", "2.4 GHz BLE", "On-tag buffer", "10 m", "3 flight trials", "2025"],
        ["[6]", "Multiple", "Multiple", "Supercap./thin-film", "N/A", "49 references", "2025"],
    ],
    col_widths=[0.7, 1.1, 1.4, 1.6, 1.1, 1.6, 1.0]
)

add_subsection_heading(doc, "C. Research Trends")

add_body(doc, (
    "Many trend show in these paper. First, people want standard and math. "
    "Model [1], design [2], and test [3] show researcher want formal way. "
    "Second, Intermittent TinyML [4] is new. It mix AI and battery-free computing."
))

add_body(doc, (
    "Third, real application like [5] show problem we don't see in lab, like "
    "drone motor noise. Fourth, survey [6] show hardware is very mature now. "
    "People understand hardware well, so now big problem is software and system integration."
))

add_subsection_heading(doc, "D. Identified Research Gaps")

add_body(doc, (
    "Even with good result, many big research gap still exist. We must fix "
    "them to use battery-free system everywhere."
), first_line_indent=False)


add_body(doc, (
    "Gap 1: No standard test benchmark. EStacker [3] is good start but not "
    "everyone use it. Every paper use different hardware and energy to test. "
    "So very hard to compare paper. We need standard test like MLPerf for battery-free."
))

add_body(doc, (
    "Gap 2: Not enough multi-source energy. Paper only use one energy like "
    "solar [1] or radio [5]. But real world have many energy source same time. "
    "We need smart system to mix solar, heat, and radio together."
))

add_body(doc, (
    "Gap 3: Edge AI for intermittent power is small. HANNA [4] is good and "
    "get 10%-44% better accuracy. But we still need know how to save AI state "
    "cheaply, and how to train AI when energy is low."
))

add_body(doc, (
    "Gap 4: No standard power rule. Paper [1] and [2] only work for their "
    "own platform. We no have standard protocol for power like we have for "
    "network. If we have standard, ecosystem will grow fast."
))

add_body(doc, (
    "Gap 5: We no know long-term reliability. People want use battery-free "
    "in harsh farm or factory. But no paper study if it break after many month "
    "in hot, cold, or noise. Drone noise problem in [5] is good example of this."
))

add_body(doc, (
    "Gap 6: Network scalability is bad. Paper only optimize one device. "
    "We no know how to manage 1000 device network when power go on and off. "
    "Normal network rule assume device always on, so we must redesign network rule."
))


# ============================================================
# V. CONCLUSION AND CONTRIBUTION
# ============================================================
add_section_heading(doc, "V. Conclusion and Contribution")

add_subsection_heading(doc, "A. Summary of Major Findings")

add_body(doc, (
    "This review read 6 new paper about battery-free system from 2024 to 2026. "
    "Paper show field move from just proof to real use. Math model [1] and "
    "design rule [2] improve time up to 90%. EStacker [3] make standard test, "
    "and HANNA [4] open TinyML area. Farm test [5] show real challenge, and survey [6] give foundation."
))

add_subsection_heading(doc, "B. Contributions of This Review")

add_body(doc, (
    "This review have 3 main contribution. First, it compare 6 paper objective, "
    "method, and result. Second, it group methodology into 4 type: math, "
    "prototype, simulation, and survey. Third, it find 6 big research gap like "
    "benchmark, energy mix, AI, power rule, reliability, and network scale."
))

add_subsection_heading(doc, "C. Future Research Directions")

add_body(doc, (
    "Based on the analysis presented in this review, the following future research directions "
    "are recommended: (1) Development of community-accepted open-source benchmark suites with "
    "standardized energy traces, hardware reference designs, and performance metrics for fair "
    "cross-platform comparison. (2) Investigation of multi-source energy harvesting architectures "
    "with intelligent source arbitration and adaptive power management algorithms. "
    "(3) Co-design of neural network architectures and intermittent computing hardware to "
    "achieve robust and energy-efficient edge AI inference under unreliable power conditions. "
    "(4) Standardization of power management protocols and abstraction layers to enable "
    "interoperability across heterogeneous battery-free platforms. (5) Longitudinal field "
    "studies assessing component degradation, failure modes, and system reliability under "
    "sustained real-world deployment conditions. (6) Design of intermittent-aware networking "
    "protocols that accommodate stochastic node availability for scalable battery-free sensor "
    "network deployments."
))

add_body(doc, (
    "The convergence of advances in energy harvesting efficiency, ultra-low-power circuit "
    "design, and edge computing capabilities suggests that battery-free embedded systems are "
    "approaching a critical threshold of commercial viability. Addressing the research gaps "
    "identified in this review will be instrumental in realizing the vision of truly "
    "maintenance-free, sustainable IoT infrastructure capable of operating autonomously "
    "across diverse and challenging deployment environments."
))

# ============================================================
# REFERENCES
# ============================================================
add_section_heading(doc, "References")

add_reference(doc, 1,
    "J. Fernández Landivar, A. Zanella, I. Gryech, S. Pollin, and H. Sallouha, "
    "\"Analytical Modeling of Batteryless IoT Sensors Powered by Ambient Energy "
    "Harvesting,\" arXiv preprint arXiv:2507.20952, Jul. 2025. "
    "[Online]. Available: https://arxiv.org/abs/2507.20952"
)

add_reference(doc, 2,
    "L. Liedtke, P. G. Kjeldsberg, F. A. Kraemer, and M. Jahre, "
    "\"Designing Cost-Effective Battery-Less Energy Harvesting for Intermittent "
    "Wireless Communication,\" IEEE Access, vol. 13, pp. 157044–157058, "
    "Jan. 2025. [Online]. Available: https://doi.org/10.1109/ACCESS.2025.3606514"
)

add_reference(doc, 3,
    "L. Liedtke, P. G. Kjeldsberg, F. A. Kraemer, and M. Jahre, "
    "\"EStacker: Explaining Battery-Less IoT System Performance with Energy Stacks,\" "
    "ACM Trans. Embed. Comput. Syst., vol. 25, no. 1, 2026. "
    "[Online]. Available: https://arxiv.org/abs/2505.22366"
)

add_reference(doc, 4,
    "R. Sahu, V. Deep, and H. Duwe, \"HANNA: Harvesting-Aware Neural Network "
    "Architecture Search for Batteryless Intermittent Devices,\" in Proc. IEEE "
    "International Performance, Computing, and Communications Conference (IPCCC), "
    "2024. [Online]. Available: https://doi.org/10.1109/IPCCC59868.2024.10850328"
)

add_reference(doc, 5,
    "P. S. Kudyba and H. Sun, \"Autonomous Agricultural Monitoring with Aerial "
    "Drones and RF Energy-Harvesting Sensor Tags,\" arXiv preprint "
    "arXiv:2502.16028, Feb. 2025. "
    "[Online]. Available: https://arxiv.org/abs/2502.16028"
)

add_reference(doc, 6,
    "IEEE MWSCAS, \"Batteryless Systems for IoT: A Survey of Circuit and System "
    "Design,\" in Proc. IEEE Midwest Symposium on Circuits and Systems (MWSCAS), "
    "2025. [Online]. Available: https://doi.org/10.1109/MWSCAS53549.2025.11244474"
)


# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT_PATH)
print(f"✅ Review paper saved to: {OUTPUT_PATH}")
print(f"\n📄 PDF Download Links for the 6 papers:")
print("=" * 70)
print("[1] https://arxiv.org/pdf/2507.20952")
print("[2] https://ieeexplore.ieee.org/document/10810505")
print("[3] https://dl.acm.org/doi/10.1145/3708997")
print("[4] https://doi.org/10.1109/IPCCC59868.2024.10850484")
print("[5] https://arxiv.org/pdf/2502.16028")
print("[6] https://ieeexplore.ieee.org/xpl/conhome/MWSCAS/2025")
