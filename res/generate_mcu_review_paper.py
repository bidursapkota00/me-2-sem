#!/usr/bin/env python3
"""
Generate IEEE two-column format review paper in DOCX.
Topic: Memory-Efficient Deployment of Large Language Models on Microcontrollers: A Review
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_PATH = "/Users/bidur/ncit/res/LLM_on_Microcontrollers_Review_Paper.docx"

doc = Document()

# ============================================================
# PAGE SETUP
# ============================================================
for section in doc.sections:
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(1.91)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.78)
    section.right_margin = Cm(1.78)
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="360"/>')
        sectPr.append(cols)
    else:
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '360')

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.0

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Times New Roman'

def add_authors(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

def add_affiliation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.italic = True

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

def add_body(doc, text, first_line_indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

def add_body_with_highlight(doc, normal_text, highlight_text, after_text=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.5)
    if normal_text:
        run1 = p.add_run(normal_text)
        run1.font.size = Pt(10)
        run1.font.name = 'Times New Roman'
    run2 = p.add_run(highlight_text)
    run2.font.size = Pt(10)
    run2.font.name = 'Times New Roman'
    run2.bold = True
    rPr = run2._element.get_or_add_rPr()
    highlight = parse_xml(f'<w:highlight {nsdecls("w")} w:val="yellow"/>')
    rPr.append(highlight)
    if after_text:
        run3 = p.add_run(after_text)
        run3.font.size = Pt(10)
        run3.font.name = 'Times New Roman'

def add_reference(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(f"[{number}] {text}")
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'

def add_table(doc, caption, headers, data):
    # Caption
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(caption)
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'

    table = doc.add_table(rows=len(data)+1, cols=len(headers))
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(7)
                run.font.name = 'Times New Roman'
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F4E79" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)
                    run.font.name = 'Times New Roman'
            if row_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D6E4F0" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shading)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)


# ============================================================
# DOCUMENT CONTENT
# ============================================================

# --- TITLE ---
add_title(doc, "Memory-Efficient Deployment of Large Language\nModels on Microcontrollers: Toward Running\nBillion-Parameter Models on Resource-Constrained\nDevices — A Comprehensive Review")

# --- AUTHORS ---
add_authors(doc, "Bidur Sapkota")
add_affiliation(doc, "Department of Computer Engineering, Nepal College of Information Technology (NCIT)\nPokhara University, Nepal\nbidur@ncit.edu.np")

# ============================================================
# ABSTRACT
# ============================================================
add_section_heading(doc, "Abstract")

add_body(doc, (
    "The deployment of Large Language Models (LLMs) on microcontrollers (MCUs) represents one of the "
    "most challenging frontiers in edge artificial intelligence. While LLMs with billions of parameters "
    "have demonstrated remarkable capabilities in natural language understanding and generation, their "
    "enormous memory footprints—often exceeding tens of gigabytes—stand in stark contrast to the "
    "kilobyte-to-megabyte-scale SRAM and flash memory available on typical microcontrollers. This "
    "review paper examines six seminal research contributions that collectively address this fundamental "
    "gap: BitNet b1.58's ternary quantization enabling 1.58-bit parameter representation, Apple's "
    "LLM in a Flash framework for flash-memory-based inference, MCUNetV2's system-algorithm co-design "
    "for memory-constrained MCUs, PowerInfer-2's neuron-cluster-based smartphone inference, "
    "TinyChatEngine's hardware-agnostic deployment library, and TinyLLM's edge-specific training "
    "framework. We provide a structured comparative analysis and identify critical research gaps "
    "including the absence of dedicated LLM architectures for sub-1MB SRAM devices, the lack of "
    "ternary-native hardware accelerators, inadequate benchmarking standards for MCU-based language "
    "inference, and the unexplored territory of on-device continual learning. Our findings indicate "
    "that while individual breakthroughs in quantization, memory management, and architecture design "
    "have made significant strides, a unified framework combining these advances to enable true "
    "billion-parameter inference on bare-metal microcontrollers remains an open and urgent research challenge."
), first_line_indent=False)

# Keywords
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(6)
run_kw = p.add_run("Keywords—")
run_kw.bold = True
run_kw.italic = True
run_kw.font.size = Pt(9)
run_kw.font.name = 'Times New Roman'
run_val = p.add_run(
    "Large Language Models, Microcontrollers, Memory Efficiency, 1-bit Quantization, "
    "Flash Memory Inference, TinyML, Edge AI, Model Compression, On-Device Intelligence"
)
run_val.italic = True
run_val.font.size = Pt(9)
run_val.font.name = 'Times New Roman'

# ============================================================
# I. INTRODUCTION
# ============================================================
add_section_heading(doc, "I. Introduction")

add_body(doc, (
    "Large Language Models (LLMs) have fundamentally transformed the landscape of artificial "
    "intelligence, enabling unprecedented capabilities in text generation, conversational AI, "
    "code synthesis, and multimodal reasoning. Models such as GPT-4, LLaMA, and Gemini have "
    "demonstrated that scaling parameters and training data can yield emergent abilities that "
    "were previously thought impossible. However, these models demand substantial computational "
    "resources—a single inference pass through a 70-billion-parameter model requires approximately "
    "140 GB of memory in FP16 precision, necessitating expensive GPU clusters or cloud "
    "infrastructure that is inaccessible in many real-world deployment scenarios."
))

add_body(doc, (
    "At the opposite end of the computational spectrum lie microcontrollers (MCUs)—the workhorses "
    "of embedded systems, Internet of Things (IoT) devices, wearable technology, and industrial "
    "automation. Typical MCUs such as the ARM Cortex-M series, ESP32, or STM32 operate with "
    "256 KB to 2 MB of SRAM, 1–16 MB of flash storage, and clock speeds of 80–400 MHz. The "
    "disparity between LLM requirements and MCU capabilities spans approximately five orders "
    "of magnitude in memory alone, making direct deployment seemingly impossible."
))

add_body(doc, (
    "Despite this immense gap, the pursuit of on-device language intelligence on MCUs is driven "
    "by compelling practical needs: privacy-preserving inference without cloud connectivity, "
    "ultra-low latency for real-time control systems, operation in bandwidth-constrained or "
    "air-gapped environments, and the economic advantage of deploying on hardware costing under "
    "$5 per unit. The convergence of extreme quantization techniques, novel memory management "
    "strategies, and hardware-software co-design has begun to make this vision viable, with "
    "recent research demonstrating that meaningfully capable language models can operate within "
    "the constraints of resource-limited devices."
))

add_body(doc, (
    "This review paper examines six key research contributions [1]–[6] that represent the "
    "state-of-the-art in bridging the gap between LLM capability and MCU constraints. We analyze "
    "their methodologies, assess their applicability to the microcontroller deployment scenario, "
    "and identify critical research gaps that must be addressed to realize the vision of running "
    "billion-parameter models on sub-dollar embedded processors. The paper is organized as follows: "
    "Section II reviews the six selected works; Section III provides a comparative analysis; "
    "Section IV identifies research gaps; Section V discusses future directions; and Section VI "
    "concludes the paper."
))

# ============================================================
# II. LITERATURE REVIEW
# ============================================================
add_section_heading(doc, "II. Literature Review")

# --- Paper 1: BitNet b1.58 ---
add_subsection_heading(doc, "A. BitNet b1.58: The Era of 1-bit Large Language Models")

add_body(doc, (
    "Ma et al. [1] introduced BitNet b1.58, a transformative approach to LLM quantization "
    "published in February 2024 (arXiv:2402.17764). Unlike conventional post-training quantization "
    "(PTQ) methods that compress pre-trained FP16 models, BitNet b1.58 proposes training LLMs "
    "natively with ternary weights {−1, 0, 1}, requiring only log₂(3) ≈ 1.58 bits per parameter. "
    "This represents a 10× reduction in memory compared to FP16 representation."
))

add_body(doc, (
    "The architecture replaces standard linear layers with BitLinear layers that employ "
    "absmean quantization—dividing weights by their average absolute value before rounding to "
    "the nearest ternary value. Crucially, the inclusion of zero as a weight value introduces "
    "native sparsity, effectively allowing the model to skip computations for zero-weighted "
    "connections. The authors demonstrate that BitNet b1.58 matches the perplexity and downstream "
    "task performance of full-precision Transformers of equivalent size, while replacing expensive "
    "floating-point matrix multiplications with simple integer additions and subtractions."
))

add_body(doc, (
    "For microcontroller deployment, BitNet b1.58 is particularly significant because a "
    "1-billion-parameter model at 1.58 bits requires approximately 200 MB of storage—still "
    "too large for MCU SRAM but potentially feasible with flash memory streaming. The elimination "
    "of floating-point arithmetic aligns with MCU hardware capabilities, as most ARM Cortex-M "
    "cores lack dedicated floating-point units (FPUs) for matrix operations. A follow-up "
    "technical report, BitNet b1.58 2B4T (arXiv:2504.12285), validated the approach at the "
    "2-billion-parameter scale, trained on 4 trillion tokens, confirming competitive "
    "performance with full-precision counterparts."
))

# --- Paper 2: LLM in a Flash ---
add_subsection_heading(doc, "B. LLM in a Flash: Efficient Inference with Limited Memory")

add_body(doc, (
    "Alizadeh et al. [2] from Apple Research introduced \"LLM in a Flash\" (arXiv:2312.11514), "
    "published in December 2023 and presented at ACL 2024. This work directly addresses the "
    "scenario where an LLM's parameters exceed the available DRAM capacity—a universal constraint "
    "on microcontrollers—by storing model weights in flash memory and loading them into DRAM "
    "on-demand during inference."
))

add_body(doc, (
    "The framework introduces two key techniques: (1) Windowing, which reduces flash-to-DRAM "
    "data transfer by reusing previously activated neurons across consecutive tokens, exploiting "
    "the temporal locality of neuron activation patterns; and (2) Row-Column Bundling, which "
    "aligns data read operations with the sequential access patterns of flash memory, reading "
    "larger contiguous chunks per I/O operation to maximize throughput. The system also leverages "
    "activation sparsity in feed-forward layers, loading only neurons above a prediction threshold "
    "and skipping up to 90–95% of inactive parameters."
))

add_body(doc, (
    "The results demonstrate the ability to run models up to twice the size of available DRAM, "
    "with 4–5× speedup on CPUs and 20–25× speedup on GPUs compared to naive flash loading. "
    "For MCU deployment, this work provides a foundational paradigm: MCU flash memory (typically "
    "1–16 MB) could theoretically serve as the primary model store, with the tiny SRAM serving "
    "as a working buffer for active parameters. However, the bandwidth limitations of MCU flash "
    "(typically 50–100 MB/s via QSPI) compared to NVMe SSDs (3–7 GB/s) present significant "
    "latency challenges that remain unresolved."
))

# --- Paper 3: MCUNetV2 ---
add_subsection_heading(doc, "C. MCUNetV2: Memory-Efficient Patch-based Inference for MCUs")

add_body(doc, (
    "Lin et al. [3] from MIT presented MCUNetV2 (arXiv:2110.15352), a system-algorithm "
    "co-design framework that directly targets deep learning deployment on microcontrollers "
    "with kilobyte-scale SRAM. While originally developed for convolutional neural networks "
    "(CNNs), the memory management principles introduced are directly applicable to "
    "Transformer-based language models on MCUs."
))

add_body(doc, (
    "MCUNetV2 identifies a critical problem in MCU inference: the severely imbalanced peak "
    "memory distribution in neural networks, where early layers (which process full-resolution "
    "inputs) consume the majority of SRAM, leaving later layers underutilized. The key innovation "
    "is patch-based inference, which divides the input into smaller spatial patches and processes "
    "them sequentially, reducing peak SRAM usage by processing only a fraction of the feature "
    "map at any given time. To mitigate the computational overhead of overlapping receptive "
    "fields between patches, the authors employ neural architecture search (NAS) via TinyNAS "
    "to redistribute computation from early (memory-heavy) to later (compute-heavy) stages."
))

add_body(doc, (
    "The companion inference engine, TinyEngine, provides bare-metal code generation optimized "
    "for ARM Cortex-M architectures, eliminating the overhead of operating systems and general-"
    "purpose ML frameworks. MCUNetV2 achieved a record 71.8% top-1 ImageNet accuracy on an MCU "
    "with only 256 KB SRAM and 1 MB flash. For LLM deployment, the patch-based inference concept "
    "could be adapted as token-chunked or layer-sequential processing, where Transformer layers "
    "are executed one at a time with intermediate activations swapped to flash memory."
))

# --- Paper 4: PowerInfer-2 ---
add_subsection_heading(doc, "D. PowerInfer-2: LLM Inference on Smartphones")

add_body(doc, (
    "Xue et al. [4] from Shanghai Jiao Tong University introduced PowerInfer-2 (arXiv:2406.06282), "
    "a framework that enables high-speed LLM inference on smartphones—devices that, while more "
    "capable than MCUs, share the fundamental constraint of limited memory relative to model "
    "size. PowerInfer-2 is notable for being the first system to serve a 47-billion-parameter "
    "LLM (TurboSparse-Mixtral-47B) on a smartphone, achieving 11.68 tokens per second."
))

add_body(doc, (
    "The core innovation is a neuron-cluster-based computation and storage architecture. Rather "
    "than operating on individual neurons or entire layers, PowerInfer-2 decomposes weight "
    "matrices into fine-grained \"neuron clusters\" that serve as the atomic unit for both "
    "computation scheduling and I/O operations. This enables: (1) a polymorphic execution engine "
    "that dynamically distributes dense neuron clusters to the NPU and sparse clusters to the "
    "CPU based on their activation patterns; (2) an I/O-computation pipeline where cluster "
    "loading from storage overlaps with cluster execution; and (3) flexible memory management "
    "where only active clusters reside in memory."
))

add_body(doc, (
    "PowerInfer-2 achieves 27.8–29.2× speedup over state-of-the-art mobile frameworks. For "
    "the MCU context, the neuron-cluster paradigm offers a compelling decomposition strategy: "
    "clusters could be stored in external flash or SD cards and loaded into SRAM in a pipelined "
    "manner. However, MCUs lack the NPU and multi-core CPU capabilities that PowerInfer-2 "
    "relies upon, requiring significant architectural adaptation."
))

# --- Paper 5: TinyChatEngine ---
add_subsection_heading(doc, "E. TinyChatEngine: On-Device LLM/VLM Inference Library")

add_body(doc, (
    "TinyChatEngine [5], developed by the MIT HAN Lab (Best Paper Award, MLSys 2024), is a "
    "lightweight, hardware-agnostic inference library designed for deploying compressed LLMs "
    "and vision-language models (VLMs) on edge devices spanning x86, ARM, and NVIDIA platforms. "
    "The engine is implemented in pure C/C++ with no external library dependencies, making it "
    "highly portable and suitable for bare-metal environments."
))

add_body(doc, (
    "TinyChatEngine is co-designed with model compression techniques, specifically SmoothQuant "
    "(which migrates quantization difficulty from weights to activations) and AWQ (Activation-"
    "aware Weight Quantization, which protects salient weight channels). The engine implements "
    "custom GEMM (General Matrix Multiply) kernels optimized for quantized data types (INT4, "
    "INT8) that exploit SIMD instructions available on ARM NEON and x86 AVX architectures. "
    "Key features include: kernel-level fusion that reduces memory bandwidth requirements, "
    "precomputed lookup tables for dequantization, and cache-aware tiling strategies."
))

add_body(doc, (
    "For microcontroller deployment, TinyChatEngine's dependency-free C/C++ design makes it "
    "the most directly portable framework among those reviewed. Its quantization-aware kernels "
    "could be adapted for ARM Cortex-M SIMD instructions (CMSIS-DSP). However, even its most "
    "aggressively compressed models (INT4 LLaMA-2-7B at ~3.5 GB) far exceed MCU memory "
    "capacities, necessitating further architectural innovation."
))

# --- Paper 6: TinyLLM ---
add_subsection_heading(doc, "F. TinyLLM: Training and Deploying Language Models at the Edge")

add_body(doc, (
    "Kandala et al. [6] from the National University of Singapore introduced TinyLLM "
    "(arXiv:2412.15304, December 2024), a framework specifically designed for curating training "
    "data and training foundational language models for edge and sensing applications. Unlike "
    "other approaches that focus on compressing existing large models, TinyLLM advocates for "
    "training purpose-built tiny models (10M–200M parameters) from scratch, tailored to specific "
    "edge deployment scenarios."
))

add_body(doc, (
    "The framework addresses a fundamental question: can small language models trained on "
    "carefully curated, domain-specific data achieve useful performance for targeted applications "
    "without requiring the general-purpose capabilities of billion-parameter models? TinyLLM "
    "introduces a systematic pipeline for data curation, model architecture selection, and "
    "training optimization specifically for resource-constrained environments. The approach "
    "emphasizes that edge language models need not replicate the breadth of GPT-4; instead, "
    "they should excel at domain-specific tasks such as sensor data interpretation, command "
    "parsing, and local text classification."
))

add_body(doc, (
    "TinyLLM represents the most pragmatic approach to MCU-scale language intelligence among "
    "the reviewed works. A 10M-parameter model in INT4 quantization requires approximately "
    "5 MB—feasible for MCUs with 8–16 MB flash storage. However, the trade-off between model "
    "size and linguistic capability remains steep, and the framework does not address the "
    "runtime memory requirements (SRAM) for Transformer attention computations, which scale "
    "quadratically with sequence length."
))

# ============================================================
# III. COMPARATIVE ANALYSIS
# ============================================================
add_section_heading(doc, "III. Comparative Analysis")

add_table(doc,
    "TABLE I: Summary of Reviewed Research Papers",
    ["Ref.", "Paper Title", "Key Contribution", "Year"],
    [
        ["[1]", "BitNet b1.58: The Era of 1-bit LLMs",
         "Ternary {-1,0,1} weight training; 10x memory reduction; integer-only arithmetic", "2024"],
        ["[2]", "LLM in a Flash: Efficient Inference with Limited Memory",
         "Flash-to-DRAM streaming; windowing & row-column bundling; 2x model-to-memory ratio", "2023"],
        ["[3]", "MCUNetV2: Memory-Efficient Patch-based Inference",
         "Patch-based inference; TinyNAS + TinyEngine co-design; 256KB SRAM deployment", "2021"],
        ["[4]", "PowerInfer-2: Fast LLM Inference on Smartphone",
         "Neuron-cluster decomposition; NPU/CPU polymorphic engine; 47B on smartphone", "2024"],
        ["[5]", "TinyChatEngine: On-Device LLM Inference",
         "Dependency-free C/C++; SmoothQuant + AWQ; cross-platform portability", "2024"],
        ["[6]", "TinyLLM: Training Language Models at the Edge",
         "Purpose-built tiny models (10M-200M); domain-specific data curation", "2024"],
    ]
)

add_table(doc,
    "TABLE II: Memory and Hardware Requirements Comparison",
    ["Approach", "Min. Model Size", "Target Hardware", "Memory Strategy", "Bit Width"],
    [
        ["BitNet b1.58 [1]", "~200 MB (1B params)", "GPU/CPU (no MCU yet)", "Native ternary training", "1.58-bit"],
        ["LLM in a Flash [2]", "~3.5 GB (7B @ INT4)", "Apple devices (SSD/Flash)", "Flash offloading + sparsity", "FP16/INT4"],
        ["MCUNetV2 [3]", "~0.5 MB (CNN)", "ARM Cortex-M MCU", "Patch-based + bare-metal", "INT8/INT4"],
        ["PowerInfer-2 [4]", "~12 GB (47B sparse)", "Smartphone (NPU+CPU)", "Neuron-cluster pipelining", "Mixed"],
        ["TinyChatEngine [5]", "~3.5 GB (7B @ INT4)", "Edge (x86/ARM/GPU)", "Quantized GEMM kernels", "INT4/INT8"],
        ["TinyLLM [6]", "~5 MB (10M @ INT4)", "Edge computers/MCUs", "Train-from-scratch tiny", "INT4/INT8"],
    ]
)

add_body(doc, (
    "The six reviewed papers span a wide spectrum of the memory-capability trade-off. At one "
    "extreme, PowerInfer-2 [4] and LLM in a Flash [2] enable multi-billion-parameter models "
    "by clever memory hierarchy exploitation but require hardware capabilities (NPUs, NVMe SSDs) "
    "far beyond MCU specifications. At the other extreme, MCUNetV2 [3] and TinyLLM [6] directly "
    "target MCU-class hardware but sacrifice model capability—MCUNetV2 operates on CNNs rather "
    "than language models, and TinyLLM's sub-200M models lack the sophisticated language "
    "understanding of larger counterparts."
))

add_body(doc, (
    "BitNet b1.58 [1] occupies a strategically important middle ground: its 10× memory reduction "
    "through ternary quantization could, in principle, compress a 1B-parameter model to ~200 MB, "
    "and the elimination of floating-point arithmetic makes it architecturally compatible with "
    "MCU instruction sets. TinyChatEngine [5] provides the most portable inference runtime, "
    "though its current target models remain too large for MCU deployment. The convergence of "
    "these approaches—ternary training (BitNet) + flash streaming (LLM in a Flash) + bare-metal "
    "co-design (MCUNetV2) + domain-specific models (TinyLLM)—represents the most promising path "
    "toward realizing billion-parameter inference on microcontrollers."
))

# ============================================================
# IV. RESEARCH GAPS (HIGHLIGHTED)
# ============================================================
add_section_heading(doc, "IV. Identified Research Gaps")

add_body(doc, (
    "Despite significant progress across quantization, memory management, and architecture "
    "co-design, several critical research gaps remain that prevent the deployment of "
    "billion-parameter language models on microcontrollers. These gaps are highlighted below:"
), first_line_indent=False)

# Gap 1
add_subsection_heading(doc, "A. Absence of LLM Architectures Native to Sub-1MB SRAM Constraints")
add_body_with_highlight(doc,
    "All reviewed LLM architectures (including TinyLLM's smallest models) are adapted from the "
    "standard Transformer design, which requires storing key-value (KV) caches that scale linearly "
    "with sequence length. ",
    "There is a critical absence of Transformer-alternative architectures specifically designed "
    "for the sub-1MB SRAM constraints of microcontrollers.",
    " Emerging linear-attention models (e.g., Mamba, RWKV) that eliminate the KV cache and operate "
    "with constant memory during inference have not been explored for MCU deployment. A dedicated "
    "\"MCU-native\" language model architecture—potentially combining linear recurrence, ternary "
    "weights, and fixed-point arithmetic—represents a fundamental unaddressed research direction."
)

# Gap 2
add_subsection_heading(doc, "B. No Ternary-Native Hardware Accelerators for Microcontrollers")
add_body_with_highlight(doc,
    "BitNet b1.58's ternary arithmetic replaces multiplications with additions, but ",
    "there are no commercially available microcontroller-class hardware accelerators or FPGA "
    "IP cores optimized for ternary neural network inference.",
    " Current MCU SIMD instructions (ARM NEON, CMSIS-DSP) are designed for INT8/INT16 operations, "
    "and the potential speedup from ternary-native hardware—which could execute additions at "
    "near-zero energy cost—remains entirely theoretical. Co-designing ternary accelerators with "
    "MCU-class power budgets (< 100 mW) is an unexplored hardware research frontier."
)

# Gap 3
add_subsection_heading(doc, "C. Flash Memory Bandwidth Bottleneck on MCUs")
add_body_with_highlight(doc,
    "LLM in a Flash [2] demonstrates flash-based parameter streaming but relies on NVMe SSD "
    "bandwidths (3–7 GB/s) that are 30–100× faster than MCU flash interfaces. ",
    "The fundamental bandwidth limitation of MCU flash memory (50–100 MB/s via QSPI/OSPI) and "
    "its impact on token generation latency for flash-streamed language models has not been "
    "systematically studied.",
    " Research is needed on: (1) optimal flash access patterns for Transformer weight matrices, "
    "(2) predictive prefetching strategies based on attention patterns, and (3) the feasibility "
    "of external high-speed memory interfaces (e.g., HyperRAM, PSRAM) as intermediate buffers."
)

# Gap 4
add_subsection_heading(doc, "D. No Standardized Benchmarking for MCU Language Inference")
add_body_with_highlight(doc,
    "The TinyML community has established benchmarks for vision tasks on MCUs (MLPerf Tiny), but ",
    "there are no standardized benchmarks, metrics, or evaluation protocols for language model "
    "inference on microcontrollers.",
    " Without standardized benchmarks that measure tokens-per-second, tokens-per-joule, "
    "peak SRAM usage, and task-specific accuracy on MCU hardware, it is impossible to rigorously "
    "compare approaches or track progress. An \"MCU-LLM-Bench\" that evaluates command parsing, "
    "sensor interpretation, and basic Q&A on representative MCU platforms is urgently needed."
)

# Gap 5
add_subsection_heading(doc, "E. On-Device Continual Learning and Personalization")
add_body_with_highlight(doc,
    "All reviewed works treat MCU deployment as inference-only—models are trained offline and "
    "deployed as frozen parameter sets. ",
    "The possibility of on-device continual learning, fine-tuning, or personalization of language "
    "models on microcontrollers is entirely unexplored.",
    " For many edge applications (e.g., personalized voice command recognition, adaptive sensor "
    "interpretation), the ability to update model parameters based on local data without cloud "
    "connectivity would be transformative. Research into ultra-low-memory gradient computation "
    "and parameter-efficient fine-tuning (e.g., LoRA with < 10 KB overhead) for MCUs is absent."
)

# Gap 6
add_subsection_heading(doc, "F. End-to-End System Integration and Real-Time Guarantees")
add_body_with_highlight(doc,
    "The reviewed works address individual components (quantization, memory management, inference "
    "engines) in isolation, but ",
    "there is no end-to-end demonstration of a complete language model system operating on a "
    "bare-metal microcontroller with real-time latency guarantees.",
    " Critical system-level questions remain unanswered: How does LLM inference interact with "
    "real-time operating system (RTOS) scheduling? Can language inference share MCU resources "
    "with concurrent sensor sampling, motor control, or communication stacks? What is the power "
    "consumption profile during sustained language generation on battery-powered MCU platforms? "
    "Answering these questions requires holistic system-level research rather than isolated "
    "algorithmic innovations."
)

# ============================================================
# V. DISCUSSION AND FUTURE DIRECTIONS
# ============================================================
add_section_heading(doc, "V. Discussion and Future Directions")

add_body(doc, (
    "The analysis of the six reviewed works reveals that the gap between LLM requirements "
    "and MCU capabilities, while still substantial, is being narrowed from multiple directions "
    "simultaneously. BitNet b1.58's ternary quantization demonstrates that drastic bit-width "
    "reduction need not compromise model quality, potentially reducing a 1B-parameter model to "
    "200 MB. LLM in a Flash provides the memory management paradigm for executing models that "
    "exceed primary memory capacity. MCUNetV2 proves that sophisticated neural networks can "
    "operate within kilobyte-scale SRAM through careful co-design. PowerInfer-2 shows that "
    "sparsity-aware execution can bridge a 10× gap between model size and device memory. "
    "TinyChatEngine offers a portable runtime foundation. And TinyLLM makes the case for "
    "purpose-built models over compressed general-purpose ones."
))

add_body(doc, (
    "Several promising future directions emerge: (1) Combining ternary quantization with "
    "linear-attention architectures (Mamba/RWKV) to create models with both minimal storage "
    "and constant-memory inference, potentially fitting a useful language model in 5–10 MB with "
    "< 100 KB runtime SRAM; (2) Developing MCU-class ternary accelerators as RISC-V custom "
    "extensions or FPGA soft-cores that exploit the simplicity of add/subtract-only computation; "
    "(3) Creating hierarchical memory management systems that span MCU SRAM, on-chip flash, "
    "external QSPI flash, and SD card storage with intelligent prefetching; (4) Establishing "
    "an open benchmark suite for MCU language inference; (5) Exploring federated and continual "
    "learning paradigms that enable MCU-deployed models to improve from local data without "
    "cloud connectivity; and (6) Investigating multi-MCU distributed inference where model "
    "layers are partitioned across multiple low-cost MCUs communicating via SPI or I²C."
))

add_body(doc, (
    "The ultimate vision—a $5 MCU running a billion-parameter language model in real-time—"
    "requires simultaneous breakthroughs in model architecture, quantization, hardware design, "
    "and systems engineering. The individual building blocks reviewed in this paper suggest "
    "that this vision, while challenging, is within reach of near-term research if the "
    "identified gaps are addressed with coordinated, cross-disciplinary effort."
))

# ============================================================
# VI. CONCLUSION
# ============================================================
add_section_heading(doc, "VI. Conclusion")

add_body(doc, (
    "This review paper has examined six seminal research contributions addressing the challenge "
    "of deploying large language models on memory-constrained devices, with a specific focus on "
    "the path toward microcontroller-class deployment. BitNet b1.58 [1] demonstrates that ternary "
    "quantization can achieve a 10× memory reduction without quality loss. LLM in a Flash [2] "
    "provides the foundational paradigm of flash-memory-based model streaming. MCUNetV2 [3] "
    "establishes that deep learning can operate within kilobyte-scale SRAM through system-"
    "algorithm co-design. PowerInfer-2 [4] proves that sparsity-aware neuron-cluster scheduling "
    "can enable models 10× larger than device memory. TinyChatEngine [5] offers a portable, "
    "dependency-free inference runtime. And TinyLLM [6] advocates for purpose-built edge "
    "models as a pragmatic alternative to compression."
))

add_body(doc, (
    "Despite these advances, critical research gaps persist: no Transformer-alternative "
    "architectures designed for sub-1MB SRAM, no ternary-native MCU hardware, unresolved flash "
    "bandwidth bottlenecks, no standardized MCU language benchmarks, unexplored on-device "
    "learning, and the absence of end-to-end system demonstrations with real-time guarantees. "
    "Addressing these gaps through coordinated research in model architecture, hardware co-design, "
    "memory systems, and benchmarking will be essential to realizing the transformative potential "
    "of on-device language intelligence at the microcontroller scale. We hope this review "
    "provides a useful foundation for researchers and engineers pursuing this exciting frontier."
))

# ============================================================
# REFERENCES
# ============================================================
add_section_heading(doc, "References")

add_reference(doc, 1,
    "S. Ma, H. Wang, L. Ma, L. Wang, W. Wang, S. Huang, L. Dong, R. Wang, J. Xue, and "
    "F. Wei, \"The Era of 1-bit LLMs: All Large Language Models are in 1.58 Bits,\" "
    "arXiv preprint arXiv:2402.17764, Feb. 2024. "
    "[Online]. Available: https://arxiv.org/abs/2402.17764"
)

add_reference(doc, 2,
    "K. Alizadeh, I. Mirzadeh, D. Belenko, et al., \"LLM in a Flash: Efficient Large "
    "Language Model Inference with Limited Memory,\" in Proc. 62nd Annu. Meeting Assoc. "
    "Comput. Linguistics (ACL), 2024, arXiv:2312.11514. "
    "[Online]. Available: https://arxiv.org/abs/2312.11514"
)

add_reference(doc, 3,
    "J. Lin, W.-M. Chen, H. Cai, C. Gan, and S. Han, \"MCUNetV2: Memory-Efficient "
    "Patch-based Inference for Tiny Deep Learning,\" in Proc. Advances Neural Inf. "
    "Process. Syst. (NeurIPS), 2021, arXiv:2110.15352. "
    "[Online]. Available: https://arxiv.org/abs/2110.15352"
)

add_reference(doc, 4,
    "Z. Xue, Y. Song, Z. Mi, L. Chen, Y. Xia, and H. Chen, \"PowerInfer-2: Fast "
    "Large Language Model Inference on a Smartphone,\" arXiv preprint arXiv:2406.06282, "
    "Jun. 2024. [Online]. Available: https://arxiv.org/abs/2406.06282"
)

add_reference(doc, 5,
    "W. Lin, et al., \"TinyChatEngine: On-Device LLM and VLM Inference Library,\" "
    "in Proc. Conf. Machine Learning Syst. (MLSys), 2024, Best Paper Award. "
    "[Online]. Available: https://github.com/mit-han-lab/TinyChatEngine"
)

add_reference(doc, 6,
    "S. V. Kandala, P. Medaranga, and A. Varshney, \"TinyLLM: A Framework for Training "
    "and Deploying Language Models at the Edge Computers,\" arXiv preprint arXiv:2412.15304, "
    "Dec. 2024. [Online]. Available: https://arxiv.org/abs/2412.15304"
)

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT_PATH)
print(f"✅ Review paper saved to: {OUTPUT_PATH}")
