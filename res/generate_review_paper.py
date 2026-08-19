#!/usr/bin/env python3
"""
Generate an IEEE two-column format review paper in DOCX format.
Topic: Recent Improvements in Large Language Models (2025-2026): A Review
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = "/Users/bidur/ncit/res/LLM_Review_Paper_2025_2026.docx"

doc = Document()

# ============================================================
# PAGE SETUP - IEEE format (Letter size, narrow margins)
# ============================================================
for section in doc.sections:
    section.page_width = Cm(21.59)   # Letter width (8.5 in)
    section.page_height = Cm(27.94)  # Letter height (11 in)
    section.top_margin = Cm(1.91)    # 0.75 in
    section.bottom_margin = Cm(2.54) # 1 in
    section.left_margin = Cm(1.78)   # 0.7 in
    section.right_margin = Cm(1.78)  # 0.7 in

    # Two-column layout
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="360"/>')
        sectPr.append(cols)
    else:
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '360')

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.0

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Times New Roman'
    return p

def add_authors(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    return p

def add_affiliation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.italic = True
    return p

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    return p

def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    return p

def add_body(doc, text, first_line_indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    return p

def add_body_with_highlight(doc, normal_text, highlight_text, after_text=""):
    """Add paragraph with highlighted (yellow background) text for research gaps."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.5)

    if normal_text:
        run1 = p.add_run(normal_text)
        run1.font.size = Pt(10)
        run1.font.name = 'Times New Roman'

    run2 = p.add_run(highlight_text)
    run2.font.size = Pt(10)
    run2.font.name = 'Times New Roman'
    run2.bold = True
    # Yellow highlight
    rPr = run2._element.get_or_add_rPr()
    highlight = parse_xml(f'<w:highlight {nsdecls("w")} w:val="yellow"/>')
    rPr.append(highlight)

    if after_text:
        run3 = p.add_run(after_text)
        run3.font.size = Pt(10)
        run3.font.name = 'Times New Roman'
    return p

def add_reference(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(f"[{number}] {text}")
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'
    return p

def add_table_paper_summary(doc):
    """Add a summary table of the 6 reviewed papers."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("TABLE I: Summary of Reviewed Research Papers")
    run.bold = True
    run.font.size = Pt(8)
    run.font.name = 'Times New Roman'

    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'

    headers = ["Ref.", "Paper Title", "Key Contribution", "Year"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(7)
                run.font.name = 'Times New Roman'
        # Dark header background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F4E79" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ["[1]", "DeepSeek-R1: Incentivizing Reasoning Capability in LLMs via RL",
         "Pure RL-based reasoning without SFT; GRPO optimization", "2025"],
        ["[2]", "Qwen3 Technical Report",
         "Unified thinking/non-thinking modes; MoE architecture; 119 languages", "2025"],
        ["[3]", "Gemma 3 Technical Report",
         "Multimodal with 128K context; interleaved local/global attention", "2025"],
        ["[4]", "FlashAttention-3: Fast and Accurate Attention",
         "Warp specialization; FP8 support; 1.5-2x speedup on H100", "2024"],
        ["[5]", "EAGLE-3: Speculative Decoding with Autoregressive Prediction Heads",
         "Bayesian-optimized layer selection; multi-token prediction", "2025"],
        ["[6]", "SAFE: Sentence-Level In-generation Attribution for RAG",
         "Real-time attribution; hallucination mitigation in RAG", "2025"],
    ]

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)
                    run.font.name = 'Times New Roman'
            # Alternating row colors
            if row_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D6E4F0" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shading)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(1.0)
        row.cells[1].width = Cm(6.0)
        row.cells[2].width = Cm(7.0)
        row.cells[3].width = Cm(1.2)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)

# ============================================================
# DOCUMENT CONTENT
# ============================================================

# --- TITLE ---
add_title(doc, "Recent Improvements in Large Language Models\n(2025–2026): A Comprehensive Review")

# --- AUTHORS ---
add_authors(doc, "Bidur Sapkota")
add_affiliation(doc, "Department of Computer Engineering, Nepal College of Information Technology (NCIT)\nPokhara University, Nepal\nbidur@ncit.edu.np")

# --- ABSTRACT ---
add_section_heading(doc, "Abstract")

abstract_text = (
    "Large Language Models (LLMs) have undergone transformative advancements in 2025–2026, "
    "shifting the paradigm from parameter scaling to inference-time reasoning, architectural "
    "efficiency, and reliability improvements. This review paper examines six seminal research "
    "contributions that represent the forefront of LLM development: DeepSeek-R1's reinforcement "
    "learning-based reasoning, Qwen3's hybrid thinking architecture, Gemma 3's efficient "
    "multimodal design, FlashAttention-3's hardware-optimized attention kernels, EAGLE-3's "
    "speculative decoding framework, and SAFE's attribution-based hallucination mitigation. "
    "We provide a structured analysis of each contribution, identify critical research gaps "
    "including the lack of standardized reasoning evaluation, limited cross-architecture "
    "transferability, energy sustainability concerns, and insufficient safety frameworks for "
    "autonomous reasoning models. Our findings suggest that future research must prioritize "
    "unified benchmarking, energy-efficient training paradigms, and robust interpretability "
    "mechanisms to ensure the responsible advancement of LLM technology."
)
add_body(doc, abstract_text, first_line_indent=False)

# --- Keywords ---
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(6)
run_kw = p.add_run("Keywords—")
run_kw.bold = True
run_kw.italic = True
run_kw.font.size = Pt(9)
run_kw.font.name = 'Times New Roman'
run_val = p.add_run("Large Language Models, Reinforcement Learning, Reasoning, Mixture-of-Experts, "
                     "Efficient Attention, Speculative Decoding, Retrieval-Augmented Generation, Research Gaps")
run_val.italic = True
run_val.font.size = Pt(9)
run_val.font.name = 'Times New Roman'

# ============================================================
# I. INTRODUCTION
# ============================================================
add_section_heading(doc, "I. Introduction")

add_body(doc, (
    "The field of Natural Language Processing (NLP) has been fundamentally reshaped by "
    "the advent and rapid evolution of Large Language Models (LLMs). Since the release of "
    "GPT-3 in 2020 and the subsequent proliferation of models like ChatGPT, LLaMA, and "
    "Gemini, LLMs have demonstrated unprecedented capabilities in text generation, reasoning, "
    "code synthesis, and multimodal understanding. However, the period spanning 2025 to 2026 "
    "marks a pivotal inflection point in LLM research, characterized by a decisive shift "
    "away from the brute-force scaling of model parameters toward more nuanced strategies "
    "focused on reasoning quality, computational efficiency, and output reliability."
))

add_body(doc, (
    "The emergence of reasoning-focused models such as DeepSeek-R1 [1] and OpenAI's o-series "
    "has demonstrated that reinforcement learning (RL) can incentivize emergent reasoning "
    "behaviors without extensive supervised fine-tuning. Simultaneously, architectural innovations "
    "like Mixture-of-Experts (MoE) in Qwen3 [2] and interleaved attention mechanisms in "
    "Gemma 3 [3] have pushed the boundaries of what smaller, more efficient models can achieve. "
    "On the systems side, advances in attention computation (FlashAttention-3 [4]), inference "
    "acceleration (EAGLE-3 [5]), and output attribution (SAFE [6]) have addressed critical "
    "bottlenecks in deploying LLMs at scale."
))

add_body(doc, (
    "This review paper provides a comprehensive examination of these six landmark contributions, "
    "analyzing their methodologies, key innovations, and experimental results. Furthermore, we "
    "identify and highlight critical research gaps that remain unaddressed, offering directions "
    "for future investigation. The remainder of this paper is organized as follows: Section II "
    "presents the reviewed literature; Section III provides a comparative analysis; Section IV "
    "identifies research gaps; Section V discusses implications; and Section VI concludes the paper."
))

# ============================================================
# II. LITERATURE REVIEW
# ============================================================
add_section_heading(doc, "II. Literature Review")

# --- Paper 1: DeepSeek-R1 ---
add_subsection_heading(doc, "A. DeepSeek-R1: Incentivizing Reasoning via Reinforcement Learning")

add_body(doc, (
    "DeepSeek-R1, introduced by DeepSeek-AI in January 2025 (arXiv:2501.12948), represents "
    "a paradigm shift in training LLMs for reasoning tasks [1]. The paper demonstrates that "
    "advanced reasoning capabilities—including self-reflection, multi-step verification, and "
    "complex mathematical problem-solving—can emerge through pure reinforcement learning, "
    "without reliance on expensive human-annotated supervised fine-tuning (SFT) data."
))

add_body(doc, (
    "The key technical innovation is Group Relative Policy Optimization (GRPO), an efficient "
    "RL technique that eliminates the need for a separate critic model, thereby significantly "
    "reducing computational overhead during training. The initial variant, DeepSeek-R1-Zero, "
    "was trained using large-scale RL directly on a base model, exhibiting emergent reasoning "
    "behaviors such as \"aha moments\" where the model spontaneously re-evaluates its approach. "
    "However, R1-Zero suffered from readability issues and language mixing. The refined "
    "DeepSeek-R1 addresses these limitations through a multi-stage training pipeline incorporating "
    "cold-start data, achieving performance competitive with OpenAI's o1 on AIME 2025 benchmarks "
    "while remaining fully open-source under the MIT License."
))

# --- Paper 2: Qwen3 ---
add_subsection_heading(doc, "B. Qwen3: Unified Hybrid Thinking Architecture")

add_body(doc, (
    "The Qwen3 Technical Report (arXiv:2505.09388), published by Alibaba's Qwen team in "
    "May 2025, presents a comprehensive family of language models ranging from 0.6B to "
    "235B parameters [2]. Qwen3's most significant contribution is the introduction of a "
    "unified \"thinking mode\" for complex reasoning tasks and a \"non-thinking mode\" for "
    "rapid, latency-sensitive responses, enabled by a dynamic thinking budget mechanism "
    "that allows users to control the depth of reasoning at inference time."
))

add_body(doc, (
    "The architecture employs both dense and Mixture-of-Expert (MoE) configurations, with "
    "the flagship Qwen3-235B-A22B utilizing 235 billion total parameters but activating "
    "only 22 billion per token through sparse expert routing. This design achieves a "
    "compelling balance between model capacity and computational cost. The training methodology "
    "encompasses a four-stage pipeline: (1) pre-training on over 36 trillion tokens across "
    "119 languages, (2) long-context extension to 128K tokens using YaRN-based positional "
    "encoding, (3) reasoning-oriented RL post-training, and (4) general-purpose alignment "
    "via direct preference optimization (DPO). Qwen3 has since been extended to multimodal "
    "(Qwen3-VL), audio (Qwen3-ASR), and omni-modal (Qwen3.5-Omni) variants."
))

# --- Paper 3: Gemma 3 ---
add_subsection_heading(doc, "C. Gemma 3: Efficient Multimodal Open-Weight Model")

add_body(doc, (
    "Google DeepMind's Gemma 3 Technical Report (arXiv:2503.19786), released in March 2025, "
    "describes a family of lightweight, open-weight multimodal models ranging from 1B to 27B "
    "parameters [3]. Gemma 3's architectural innovation lies in its interleaved local/global "
    "attention mechanism, where local sliding window attention handles short-range dependencies "
    "while global attention layers capture long-range context, enabling support for 128K "
    "token context windows without the linear memory growth typical of standard full attention."
))

add_body(doc, (
    "The vision component integrates a SigLIP-based image encoder through a Pan-and-Scan "
    "strategy that adaptively crops high-resolution images into multiple sub-images, enabling "
    "fine-grained visual understanding. Gemma 3-27B achieves competitive performance with "
    "models up to twice its parameter count on benchmarks including MMLU (75.6%), GSM8K "
    "(89.0%), and HumanEval (77.4%). The model family is designed for single-GPU deployment, "
    "democratizing access to capable multimodal AI for researchers and developers with "
    "limited computational resources."
))

# --- Paper 4: FlashAttention-3 ---
add_subsection_heading(doc, "D. FlashAttention-3: Hardware-Aware Efficient Attention")

add_body(doc, (
    "FlashAttention-3 (arXiv:2407.08608), authored by Shah et al., introduces a suite of "
    "hardware-specific optimizations for attention computation on NVIDIA Hopper GPUs (H100) "
    "[4]. Building upon the I/O-aware foundations of FlashAttention-1 and FlashAttention-2, "
    "the third iteration addresses the observation that prior versions achieved only ~35% "
    "hardware utilization on H100 GPUs due to underutilization of new hardware features."
))

add_body(doc, (
    "Three core innovations drive the performance gains: (1) Warp Specialization, which "
    "exploits the asynchronous nature of Tensor Cores and the Tensor Memory Accelerator "
    "(TMA) to overlap data movement with computation; (2) Interleaved GEMM-Softmax "
    "scheduling, which pipelines block-wise matrix multiplication and softmax operations "
    "to hide computational latency; and (3) native FP8 support with block quantization "
    "and incoherent processing to maintain numerical accuracy at reduced precision. "
    "FlashAttention-3 achieves up to 740 TFLOPs/s (75% utilization) in FP16 and 1.2 "
    "PFLOPs/s in FP8, representing a 1.5–2.0× speedup over FlashAttention-2. The work "
    "has since been succeeded by FlashAttention-4, targeting NVIDIA Blackwell (B200) GPUs."
))

# --- Paper 5: EAGLE-3 ---
add_subsection_heading(doc, "E. EAGLE-3: Advanced Speculative Decoding")

add_body(doc, (
    "EAGLE-3, developed as a state-of-the-art speculative decoding framework in 2025, "
    "advances the paradigm of draft-then-verify inference acceleration for autoregressive "
    "LLMs [5]. Unlike traditional speculative decoding that requires a separate, smaller "
    "draft model, EAGLE-3 attaches autoregressive prediction heads to the target model's "
    "internal layers, enabling self-speculative inference without an external drafter."
))

add_body(doc, (
    "The key innovations include: (1) Bayesian optimization for selecting the optimal "
    "internal layer from which to derive draft predictions, balancing draft quality against "
    "computational overhead; (2) multi-token prediction capability that increases draft "
    "depth without proportional increase in compute cost; and (3) cross-attention-based "
    "decoding variants (Beagle) that simplify the architecture by eliminating auxiliary "
    "components like pooling or fusion layers. EAGLE-3 consistently achieves 2–3× speedups "
    "for LLM inference while maintaining lossless output quality, matching or exceeding "
    "the performance of EAGLE-v2 with improved training efficiency."
))

# --- Paper 6: SAFE ---
add_subsection_heading(doc, "F. SAFE: Sentence-Level Attribution for RAG Reliability")

add_body(doc, (
    "The SAFE framework (arXiv:2505.12621), proposed by Batista et al. in May 2025, "
    "addresses one of the most critical challenges in deployed LLM systems: hallucination "
    "in Retrieval-Augmented Generation (RAG) pipelines [6]. SAFE introduces a two-step "
    "in-generation attribution mechanism that (1) predicts the required number of source "
    "references for each generated sentence, and (2) performs fine-grained attribution "
    "linking generated claims to specific segments in retrieved source documents."
))

add_body(doc, (
    "Unlike post-hoc fact-checking approaches, SAFE operates during the generation process "
    "itself, enabling real-time verification and significantly reducing the latency overhead "
    "of attribution. The framework is model-agnostic and can be integrated into existing "
    "RAG pipelines with minimal architectural changes. The authors provide both the framework "
    "implementation and the training dataset publicly, facilitating reproducibility. SAFE "
    "demonstrates meaningful improvements in attribution accuracy and user trust in "
    "enterprise RAG deployments where factual reliability is paramount."
))

# ============================================================
# III. COMPARATIVE ANALYSIS
# ============================================================
add_section_heading(doc, "III. Comparative Analysis")

add_table_paper_summary(doc)

add_body(doc, (
    "The six reviewed papers collectively represent three major thrusts in contemporary "
    "LLM research: (1) reasoning and training methodology (DeepSeek-R1, Qwen3), (2) "
    "architectural and systems-level efficiency (Gemma 3, FlashAttention-3, EAGLE-3), and "
    "(3) output reliability and safety (SAFE). A notable convergence is the shift toward "
    "inference-time scaling—rather than simply training larger models, researchers are "
    "investing in techniques that allow models to \"think harder\" (DeepSeek-R1's RL-based "
    "reasoning), \"think more efficiently\" (FlashAttention-3's hardware-aware kernels, "
    "EAGLE-3's speculative decoding), or \"think more reliably\" (SAFE's attribution)."
))

add_body(doc, (
    "Another significant trend is the democratization of capabilities. Gemma 3 and Qwen3 "
    "both emphasize single-GPU deployability and open-weight distribution, making "
    "frontier-level performance accessible to researchers beyond major corporations. The "
    "MoE architecture in Qwen3 and the interleaved attention in Gemma 3 represent different "
    "but complementary strategies for achieving this goal—sparse expert activation versus "
    "memory-efficient attention patterns."
))

# ============================================================
# IV. RESEARCH GAPS (HIGHLIGHTED)
# ============================================================
add_section_heading(doc, "IV. Identified Research Gaps")

add_body(doc, (
    "Despite the substantial progress documented in the reviewed literature, several "
    "critical research gaps remain that warrant urgent investigation. These gaps are "
    "highlighted below:"
), first_line_indent=False)

# Gap 1
add_subsection_heading(doc, "A. Lack of Standardized Reasoning Evaluation Benchmarks")
add_body_with_highlight(doc,
    "While DeepSeek-R1 and Qwen3 demonstrate impressive reasoning capabilities, ",
    "there is no universally accepted benchmark suite for evaluating reasoning quality, "
    "consistency, and robustness across different LLMs.",
    " Current evaluations rely on competition-level math problems (AIME, MATH) and coding "
    "benchmarks (HumanEval, SWE-bench), which do not capture the full spectrum of reasoning "
    "abilities required in real-world applications such as legal analysis, medical diagnosis, "
    "or multi-hop scientific reasoning."
)

# Gap 2
add_subsection_heading(doc, "B. Limited Cross-Architecture Transferability of Optimizations")
add_body_with_highlight(doc,
    "FlashAttention-3's optimizations are tightly coupled to NVIDIA Hopper GPU architecture, and ",
    "there is insufficient research on transferring hardware-aware attention optimizations across "
    "heterogeneous hardware platforms",
    " (e.g., AMD MI300X, Google TPUs, Apple Silicon, Qualcomm NPUs). As edge and on-device "
    "deployment becomes increasingly important, hardware-agnostic efficiency techniques represent "
    "a critical unmet need."
)

# Gap 3
add_subsection_heading(doc, "C. Energy Sustainability and Environmental Impact")
add_body_with_highlight(doc,
    "None of the six reviewed papers comprehensively address the environmental cost of training "
    "and deploying LLMs. ",
    "There is a significant research gap in developing standardized carbon footprint metrics "
    "and energy-efficient training paradigms for LLMs.",
    " While MoE architectures and efficient attention reduce per-inference costs, the cumulative "
    "energy consumption of RL-based training pipelines (as in DeepSeek-R1) and large-scale "
    "pre-training (as in Qwen3's 36T token corpus) remains largely unreported and unoptimized."
)

# Gap 4
add_subsection_heading(doc, "D. Safety Frameworks for Autonomous Reasoning Models")
add_body_with_highlight(doc,
    "DeepSeek-R1 demonstrates that reasoning capabilities can emerge spontaneously through RL, but ",
    "there are no robust frameworks to ensure that emergent reasoning behaviors remain aligned, "
    "interpretable, and controllable.",
    " The risk of models developing opaque reasoning strategies that produce correct outputs "
    "through unintended or unsafe intermediate steps is largely unexamined. SAFE addresses "
    "attribution for RAG but does not extend to the safety of internal reasoning chains."
)

# Gap 5
add_subsection_heading(doc, "E. Speculative Decoding Under Real-World Workloads")
add_body_with_highlight(doc,
    "EAGLE-3's speculative decoding achieves impressive speedups in controlled settings, yet ",
    "the fragility of speculative decoding under diverse, real-world production workloads "
    "with varying sequence lengths, batch sizes, and request patterns is insufficiently studied.",
    " Adaptive mechanisms like Online Speculative Decoding (OSD) represent early steps, but "
    "comprehensive frameworks for dynamic draft strategy selection in production environments "
    "remain absent."
)

# Gap 6
add_subsection_heading(doc, "F. Multilingual and Low-Resource Language Reasoning")
add_body_with_highlight(doc,
    "Although Qwen3 supports 119 languages, ",
    "the reasoning capabilities of LLMs in low-resource languages remain severely underexplored.",
    " Most reasoning benchmarks and training data are heavily English-centric, creating a "
    "widening capability gap between high-resource and low-resource languages. Research on "
    "cross-lingual reasoning transfer and culturally-aware evaluation frameworks is urgently needed."
)

# ============================================================
# V. DISCUSSION AND FUTURE DIRECTIONS
# ============================================================
add_section_heading(doc, "V. Discussion and Future Directions")

add_body(doc, (
    "The research landscape of LLMs in 2025–2026 reveals a maturing field that is moving "
    "beyond the \"scaling law\" paradigm toward more sophisticated approaches to intelligence. "
    "The convergence of reinforcement learning, sparse architectures, hardware co-design, "
    "and reliability engineering suggests that the next generation of LLMs will be defined "
    "not by their parameter count but by their reasoning depth, efficiency, and trustworthiness."
))

add_body(doc, (
    "Several promising future directions emerge from this review: (1) Unified reasoning "
    "benchmarks that span multiple domains and languages; (2) Hardware-agnostic efficiency "
    "frameworks that enable deployment across diverse compute environments; (3) Carbon-aware "
    "training schedulers that optimize model quality under energy constraints; (4) Formal "
    "verification methods for reasoning chain safety; (5) Cross-lingual reasoning transfer "
    "learning that leverages high-resource language capabilities to bootstrap low-resource "
    "language reasoning; and (6) Integration of attribution mechanisms into the reasoning "
    "process itself, rather than treating generation and verification as separate concerns."
))

add_body(doc, (
    "Furthermore, the open-weight movement exemplified by DeepSeek-R1, Qwen3, and Gemma 3 "
    "has profound implications for the democratization of AI research. As these models "
    "become increasingly capable and accessible, the research community must develop "
    "governance frameworks that balance openness with safety, ensuring that the benefits "
    "of advanced reasoning models are broadly distributed while mitigating potential risks."
))

# ============================================================
# VI. CONCLUSION
# ============================================================
add_section_heading(doc, "VI. Conclusion")

add_body(doc, (
    "This review paper has examined six seminal research contributions that define the "
    "state-of-the-art in Large Language Model development during 2025–2026. From "
    "DeepSeek-R1's demonstration that pure reinforcement learning can produce emergent "
    "reasoning, to Qwen3's scalable hybrid thinking architecture, Gemma 3's efficient "
    "multimodal design, FlashAttention-3's hardware-level optimizations, EAGLE-3's "
    "inference acceleration, and SAFE's real-time attribution framework, these works "
    "collectively represent a paradigm shift toward more capable, efficient, and reliable "
    "language models."
))

add_body(doc, (
    "However, significant research gaps persist, particularly in standardized reasoning "
    "evaluation, cross-hardware optimization transferability, energy sustainability, "
    "safety of emergent reasoning behaviors, production-grade speculative decoding, and "
    "multilingual reasoning equity. Addressing these gaps will be essential for the "
    "responsible and inclusive advancement of LLM technology. We hope this review serves "
    "as a useful reference for researchers and practitioners working at the frontier of "
    "language model development."
))

# ============================================================
# REFERENCES
# ============================================================
add_section_heading(doc, "References")

add_reference(doc, 1,
    "DeepSeek-AI, \"DeepSeek-R1: Incentivizing Reasoning Capability in LLMs via "
    "Reinforcement Learning,\" arXiv preprint arXiv:2501.12948, Jan. 2025. "
    "[Online]. Available: https://arxiv.org/abs/2501.12948"
)

add_reference(doc, 2,
    "Qwen Team, \"Qwen3 Technical Report,\" arXiv preprint arXiv:2505.09388, "
    "May 2025. [Online]. Available: https://arxiv.org/abs/2505.09388"
)

add_reference(doc, 3,
    "Google DeepMind, \"Gemma 3 Technical Report,\" arXiv preprint arXiv:2503.19786, "
    "Mar. 2025. [Online]. Available: https://arxiv.org/abs/2503.19786"
)

add_reference(doc, 4,
    "J. Shah, G. Bikshandi, Y. Zhang, V. Thakkar, P. Ramani, and T. Dao, "
    "\"FlashAttention-3: Fast and Accurate Attention with Asynchrony and Low-precision,\" "
    "arXiv preprint arXiv:2407.08608, Jul. 2024. "
    "[Online]. Available: https://arxiv.org/abs/2407.08608"
)

add_reference(doc, 5,
    "Y. Li et al., \"EAGLE-3: Speculative Decoding with Autoregressive Prediction Heads,\" "
    "arXiv preprint, 2025. [Online]. Available: https://arxiv.org/abs/2503.01840"
)

add_reference(doc, 6,
    "J. E. Batista, E. Vatai, and M. Wahib, \"SAFE: Improving LLM Systems using "
    "Sentence-Level In-generation Attribution,\" arXiv preprint arXiv:2505.12621, "
    "May 2025. [Online]. Available: https://arxiv.org/abs/2505.12621"
)

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT_PATH)
print(f"✅ Review paper saved to: {OUTPUT_PATH}")
