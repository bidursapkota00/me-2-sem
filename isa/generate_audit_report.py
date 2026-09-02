"""
Generate IS Audit Report for Yeti Airlines Nepal in DOCX format.
Black & white, Times New Roman, formal academic style.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


BLACK = RGBColor(0x00, 0x00, 0x00)


def set_cell_borders(cell):
    """Set thin black borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def remove_cell_borders(cell):
    """Remove all borders from a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def make_plain_table(table):
    """Make a table plain black-bordered with no shading."""
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            # Remove any shading
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for shd in tcPr.findall(f'{{{tcPr.nsmap["w"]}}}shd'):
                tcPr.remove(shd)
            # Set font for all runs
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.color.rgb = BLACK


def style_heading(heading):
    """Make heading black Times New Roman."""
    for run in heading.runs:
        run.font.color.rgb = BLACK
        run.font.name = 'Times New Roman'


def create_document():
    doc = Document()

    # --- Document styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = BLACK

    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.5

    # Style headings to be black Times New Roman and add space after
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = 'Times New Roman'
        h_style.font.color.rgb = BLACK
        h_style.paragraph_format.space_after = Pt(12)

    # Configure section margins
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Add page numbers in footer
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add PAGE field
    run = footer_para.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = footer_para.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(10)
    run3 = footer_para.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)

    # ==================== COVER PAGE ====================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Information System Audit Report')
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = BLACK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('DRAFT REPORT')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = BLACK

    doc.add_paragraph()

    org_name = doc.add_paragraph()
    org_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org_name.add_run('Yeti Airlines Pvt. Ltd.')
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = BLACK

    doc.add_paragraph()

    # Submitted info table (borderless)
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ('Submitted By:', 'Bidur Sapkota (ME_252952)'),
        ('Submitted To:', 'Yeti Airlines Pvt. Ltd.'),
        ('Date:', 'September 2, 2026'),
        ('Version:', '1.0 (Draft)'),
    ]
    for i, (label, value) in enumerate(info_data):
        cell_label = info_table.cell(i, 0)
        cell_value = info_table.cell(i, 1)
        run_l = cell_label.paragraphs[0].add_run(label)
        run_l.bold = True
        run_l.font.size = Pt(12)
        run_l.font.name = 'Times New Roman'
        run_l.font.color.rgb = BLACK
        cell_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_v = cell_value.paragraphs[0].add_run(value)
        run_v.font.size = Pt(12)
        run_v.font.name = 'Times New Roman'
        run_v.font.color.rgb = BLACK
        cell_value.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    for row in info_table.rows:
        for cell in row.cells:
            remove_cell_borders(cell)

    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    h = doc.add_heading('TABLE OF CONTENTS', level=1)
    style_heading(h)

    toc_items = [
        ('Introduction', ''),
        ('Objectives', ''),
        ('Scope of the Work', ''),
        ('A. Information System Review Based on ISO 27001:2022 Framework', ''),
        ('    A.1. Organizational Controls', ''),
        ('        A.1.1. Policies for Information Security (A.5.1)', ''),
        ('        A.1.2. Information Security Roles and Responsibilities (A.5.2)', ''),
        ('    A.2. People Controls', ''),
        ('        A.2.1. Information Security Awareness, Education and Training (A.6.3)', ''),
        ('    A.3. Physical Controls', ''),
        ('        A.3.1. Physical Security Perimeters (A.7.1)', ''),
        ('        A.3.2. Securing Offices, Rooms and Facilities (A.7.3)', ''),
        ('    A.4. Technological Controls', ''),
        ('        A.4.1. Access Control (A.8.2)', ''),
        ('        A.4.2. Protection Against Malware (A.8.7)', ''),
        ('        A.4.3. Management of Technical Vulnerabilities (A.8.8)', ''),
        ('        A.4.4. Information Backup (A.8.13)', ''),
        ('        A.4.5. Logging (A.8.15)', ''),
        ('B. Summary of Findings', ''),
        ('C. Recommendations', ''),
        ('D. Regulatory Compliance Context', ''),
    ]

    for item, _ in toc_items:
        p = doc.add_paragraph()
        indent_level = len(item) - len(item.lstrip())
        p.paragraph_format.left_indent = Cm(indent_level * 0.3)
        run = p.add_run(item.strip())
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = BLACK
        if indent_level == 0:
            run.bold = True
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ==================== LIST OF ABBREVIATIONS ====================
    h = doc.add_heading('LIST OF ABBREVIATIONS', level=1)
    style_heading(h)

    abbreviations = [
        ('AAA', 'Authentication, Authorization and Accounting'),
        ('ATS', 'Auto Transfer Switch'),
        ('BCP', 'Business Continuity Planning'),
        ('CAAN', 'Civil Aviation Authority of Nepal'),
        ('CIA', 'Confidentiality, Integrity and Availability'),
        ('DC', 'Data Center'),
        ('DRC', 'Disaster Recovery Center'),
        ('EDR', 'Endpoint Detection and Response'),
        ('ETA', 'Electronic Transactions Act'),
        ('ICAO', 'International Civil Aviation Organization'),
        ('ICT', 'Information and Communication Technologies'),
        ('IDS', 'Intrusion Detection System'),
        ('IPS', 'Intrusion Prevention System'),
        ('IS', 'Information System'),
        ('ISMS', 'Information Security Management System'),
        ('ISO', 'International Organization for Standardization'),
        ('IT', 'Information Technology'),
        ('IVR', 'Interactive Voice Response'),
        ('MFA', 'Multi-Factor Authentication'),
        ('NTA', 'Nepal Telecommunications Authority'),
        ('PAM', 'Privilege Access Management'),
        ('RBAC', 'Role-Based Access Control'),
        ('RPO', 'Recovery Point Objective'),
        ('RTO', 'Recovery Time Objective'),
        ('SIEM', 'Security Information and Event Management'),
        ('STOL', 'Short Take-Off and Landing'),
        ('UAT', 'User Acceptance Test'),
        ('UPS', 'Uninterruptible Power Supply'),
    ]

    abbr_table = doc.add_table(rows=len(abbreviations) + 1, cols=2)
    abbr_table.style = 'Table Grid'

    hdr = abbr_table.rows[0]
    for j, text in enumerate(['Abbreviation', 'Full Form']):
        cell = hdr.cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = BLACK

    for i, (abbr, full) in enumerate(abbreviations, 1):
        for j, text in enumerate([abbr, full]):
            cell = abbr_table.cell(i, j)
            cell.text = ''
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = BLACK

    doc.add_page_break()

    # ==================== INTRODUCTION ====================
    h = doc.add_heading('Introduction', level=1)
    style_heading(h)

    doc.add_paragraph(
        'Information and Communication Technologies (ICT) plays a vital role for Yeti Airlines Pvt. Ltd. '
        'to enable its entire business process, including flight operations management, online booking systems, '
        'passenger data management, IP telephony, and ground operations coordination across its domestic network. '
        'Yeti Airlines, established in 1998 and headquartered at Tribhuvan International Airport, Kathmandu, is '
        'the second-largest domestic carrier in Nepal by passenger volume, operating a fleet of seven ATR 72-500 '
        'aircraft with scheduled services to major cities including Pokhara, Nepalgunj, Bhairahawa, Biratnagar, '
        'Bhadrapur, Janakpur, and Simara. The airline is also the parent company of Tara Air, handling STOL '
        'operations to remote mountainous regions.'
    )

    doc.add_paragraph(
        'An IS Audit enables Yeti Airlines to identify existing risks exposed from ICT services. It enables '
        'in-depth analysis of the existing technical environment, including computer applications, hardware '
        'infrastructure, IT plans, policies, adaptation of new technologies, and IT-related personnel. Hence, '
        'the IS Audit provides a comprehensive analysis of how well the technology infrastructure aligns with '
        'the goals and needs of the organization.'
    )

    doc.add_paragraph(
        'This report presents the detailed activities conducted by certified professionals in the domain of '
        'IS Audit for Yeti Airlines Pvt. Ltd. based on the Information Security Management Framework '
        '(ISO 27001:2022). It aids in preserving the Confidentiality, Integrity and Availability (CIA) of '
        'information by applying a risk management process and gives confidence to interested stakeholders '
        'that risks are adequately managed.'
    )

    # ==================== OBJECTIVES ====================
    h = doc.add_heading('Objectives', level=1)
    style_heading(h)

    doc.add_paragraph(
        'The main objective of this IS Audit is to evaluate the state of the current information system '
        'present at Yeti Airlines Pvt. Ltd., which will help mitigate existing Cyber Security Risks. '
        'The specific objectives include:'
    )

    objectives = [
        'To assess the effectiveness of information security policies and procedures in alignment with ISO 27001:2022.',
        'To evaluate the physical and environmental security controls at the Data Center (DC) and Disaster Recovery Center (DRC).',
        'To review access control mechanisms and identity management practices across critical IT systems.',
        'To assess compliance with applicable regulatory requirements including the Civil Aviation Act (1996), '
        'Electronic Transactions Act (2008), and Individual Privacy Act (2018).',
        'To identify vulnerabilities in the IT infrastructure and recommend remedial actions.',
        'To evaluate business continuity and disaster recovery preparedness.',
    ]
    for obj in objectives:
        p = doc.add_paragraph(obj, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)

    # ==================== SCOPE OF WORK ====================
    h = doc.add_heading('Scope of the Work', level=1)
    style_heading(h)

    doc.add_paragraph(
        'The scope of the IS Audit for Yeti Airlines Pvt. Ltd. is based on benchmarks provisioned by the '
        'ISO 27001:2022 framework. In order to carry out the IS Audit, we have identified gaps in current '
        'work procedures, practices, and policies. The following activities were carried out during the '
        'course of the IS Audit:'
    )

    scope_items = [
        'Review of the Information Security Policy and related documentation.',
        'Assessment of information security roles, responsibilities, and organizational structure.',
        'Evaluation of personnel security awareness and training programs.',
        'Physical security assessment of the Data Center (DC), Disaster Recovery Center (DRC), and office premises.',
        'Review of logical access controls for critical applications including flight booking systems, '
        'passenger management systems, and IVR/telephony systems.',
        'Assessment of malware protection, vulnerability management, and patch management practices.',
        'Review of backup and recovery procedures.',
        'Evaluation of logging and monitoring capabilities.',
        'Assessment of compliance with CAAN regulations and Nepal\'s Electronic Transactions Act (2008).',
        'Review of incident management and response procedures.',
    ]
    for item in scope_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph(
        'The detailed Control Objectives based on ISO 27001:2022 that were checked during the IS Audit '
        'are discussed below.'
    )

    doc.add_page_break()

    # ==================== SECTION A: IS REVIEW ====================
    h = doc.add_heading('A. Information System Review Based on ISO 27001:2022 Framework', level=1)
    style_heading(h)

    doc.add_paragraph(
        'The review of the Information Security Management System is done based on ISO 27001:2022. The '
        'information security management system preserves the confidentiality, integrity and availability '
        'of information by applying a risk management process and gives confidence to interested parties '
        'that risks are adequately managed. It is important that the information security management system '
        'is part of and integrated with Yeti Airlines\' processes and overall management structure, and that '
        'information security is considered in the design of processes, information systems, and controls.'
    )

    doc.add_paragraph(
        'The Information System Audit is conducted based on the Control Objectives of the ISO 27001:2022 '
        'Information Security Management Framework. Ten (10) key controls have been selected for evaluation. '
        'The status of each control objective is validated, and recommendations are provided during the '
        'IS Audit of Yeti Airlines Pvt. Ltd. The details of each Control Objective are given below.'
    )

    # ==================== CONTROL DEFINITIONS ====================
    controls = [
        {
            'section': 'A.1. Organizational Controls',
            'subsection': 'A.1.1. Policies for Information Security (A.5.1)',
            'control': (
                'Information security policy and topic-specific policies should be defined, approved by '
                'management, published, communicated to and acknowledged by relevant personnel and '
                'relevant interested parties, and reviewed at planned intervals and if significant changes occur.'
            ),
            'purpose': (
                'To ensure continuing suitability, adequacy, effectiveness of management direction and support '
                'for information security in accordance with business, legal, statutory, regulatory and '
                'contractual requirements.'
            ),
            'observation': (
                'Yeti Airlines has established an Information Security Policy (Version 1.0, 2021); however, '
                'the policy has not been reviewed or updated since its initial release in 2021. Given the '
                'significant organizational changes\u2014including the acquisition of 49% stake by Asian Life Insurance '
                'Co. Ltd. in late 2023, the preparation for IPO, and the post-accident regulatory reforms mandated '
                'by CAAN following the January 2023 incident\u2014the policy does not reflect the current operational '
                'reality. Additionally, no topic-specific policies (e.g., acceptable use policy, remote access '
                'policy, data classification policy) were found to be formally documented. Staff awareness of the '
                'existing policy was observed to be limited, particularly among ground crew at remote stations.'
            ),
            'risk_rating': 'HIGH',
            'recommendation': (
                'Yeti Airlines should immediately review and update its Information Security Policy to align with '
                'ISO 27001:2022 requirements and the current organizational context. The policy should be approved '
                'by top management, communicated to all employees across all stations, and acknowledged in writing. '
                'Topic-specific policies should be developed covering areas such as acceptable use, data '
                'classification, remote access, and mobile device management. A formal review cycle of at least '
                'annual review should be established, with additional reviews triggered by significant changes such '
                'as regulatory updates from CAAN or major organizational restructuring.'
            ),
            'management_response': (
                'The Information Security Policy shall be reviewed and updated by the end of fiscal year 2083/84. '
                'A dedicated working committee will be formed to draft topic-specific policies.'
            ),
        },
        {
            'subsection': 'A.1.2. Information Security Roles and Responsibilities (A.5.2)',
            'control': (
                'Information security roles and responsibilities should be defined and allocated according to '
                'the organization\'s needs.'
            ),
            'purpose': (
                'To establish a defined, approved, and understood structure for the implementation, operation '
                'and management of information security within the organization.'
            ),
            'observation': (
                'Information security roles and responsibilities at Yeti Airlines are not properly defined. The '
                'IT department handles security responsibilities in an ad-hoc manner without clear role delineation. '
                'There is no designated Chief Information Security Officer (CISO) or equivalent position. The IT '
                'Manager currently oversees both IT operations and security, leading to potential conflicts of '
                'interest. Security responsibilities at branch offices and outstations are not formally assigned. '
                'The airline operates across multiple domestic destinations, and there is no clear chain of command '
                'for information security incident escalation from remote stations to the head office in Kathmandu.'
            ),
            'risk_rating': 'HIGH',
            'recommendation': (
                'Yeti Airlines should define and formally document all information security roles and responsibilities '
                'in accordance with the information security policy. A dedicated CISO or Information Security Manager '
                'position should be created, reporting directly to senior management. Security focal points should be '
                'designated at each major station (Pokhara, Nepalgunj, Bhairahawa, Biratnagar, Bhadrapur). A clear '
                'escalation matrix for security incidents should be developed and communicated to all staff. Role-Based '
                'Access Control (RBAC) principles should be applied to all IT systems.'
            ),
            'management_response': 'Management will formally define the roles and responsibilities by Q1 of the next fiscal year. A new Information Security Manager role will be created to oversee these functions across all stations.',
        },
        {
            'section': 'A.2. People Controls',
            'subsection': 'A.2.1. Information Security Awareness, Education and Training (A.6.3)',
            'control': (
                'Personnel of the organization and relevant interested parties should receive appropriate information '
                'security awareness, education and training and regular updates of the organization\'s information '
                'security policy, topic-specific policies and procedures, as relevant for their job function.'
            ),
            'purpose': (
                'To ensure personnel and relevant interested parties are aware of and fulfill their information '
                'security responsibilities.'
            ),
            'observation': (
                'Yeti Airlines does not have a structured information security awareness and training program. '
                'While the airline has invested in safety training programs aligned with CAAN regulations and IATA '
                'ISSA certification requirements, cybersecurity awareness training has not been formalized. '
                'Interviews with staff revealed that employees are not aware of basic cybersecurity practices such as '
                'phishing identification, password management, and social engineering threats. No records of '
                'cybersecurity training sessions were available for review. The online booking platform and customer '
                'service systems (IVR/IP telephony) handle sensitive passenger data, making staff awareness critically '
                'important. The airline\'s ground handling staff at multiple stations were found to share login '
                'credentials for convenience.'
            ),
            'risk_rating': 'HIGH',
            'recommendation': (
                'Yeti Airlines should establish a comprehensive information security awareness and training program '
                'that covers all employees, including flight crew, ground staff, and management. Training should be '
                'conducted at planned intervals (at least annually) and should cover topics such as phishing awareness, '
                'password hygiene, social engineering, data handling procedures, and incident reporting. Specialized '
                'training should be provided to IT staff on secure coding, system administration security, and incident '
                'response. Training records should be maintained, and effectiveness should be assessed through periodic '
                'tests and simulated phishing exercises.'
            ),
            'management_response': 'A comprehensive information security awareness training program will be initiated in the upcoming quarter. Mandatory annual training sessions will be implemented for all employees, and records will be maintained by the HR department.',
        },
        {
            'section': 'A.3. Physical Controls',
            'subsection': 'A.3.1. Physical Security Perimeters (A.7.1)',
            'control': (
                'Security perimeters should be defined and used to protect areas that contain information and '
                'other associated assets.'
            ),
            'purpose': (
                'To prevent unauthorized physical access, damage and interference to the organization\'s '
                'information and other associated assets.'
            ),
            'observation': (
                'The Data Center (DC) located at the Yeti Airlines head office in Kathmandu has basic physical '
                'security measures in place, including locked doors and restricted entry. However, the following '
                'gaps were identified: (1) Access to the DC is controlled through a simple key-lock mechanism '
                'rather than electronic access control systems (biometric or card-based). (2) No visitor log is '
                'maintained for the DC area. (3) The Disaster Recovery Center (DRC) physical security was found '
                'to be inadequate, with shared access corridors and no dedicated security perimeter. (4) CCTV '
                'coverage at the DC entrance was found to be limited, with footage retained for only 15 days. '
                '(5) Multiple branch offices across Nepal were observed to have minimal physical security for '
                'their local IT equipment and networking infrastructure.'
            ),
            'risk_rating': 'HIGH',
            'recommendation': (
                'Yeti Airlines should implement electronic access control systems (biometric or smart card-based) '
                'for the Data Center and Disaster Recovery Center. A visitor management system with proper logging '
                'should be deployed. CCTV coverage should be enhanced with a minimum retention period of 90 days '
                'as per industry best practices. The DRC should have a dedicated, well-defined security perimeter '
                'with independent access controls. Physical security assessments should be conducted at all branch '
                'offices, and minimum security standards should be established for IT equipment rooms at all '
                'stations.'
            ),
            'management_response': 'Budget has been allocated for upgrading physical security measures. Electronic access control systems and enhanced CCTV coverage for the Data Center and Disaster Recovery Center will be deployed within six months.',
        },
        {
            'subsection': 'A.3.2. Securing Offices, Rooms and Facilities (A.7.3)',
            'control': (
                'Physical security for offices, rooms and facilities should be designed and implemented.'
            ),
            'purpose': (
                'To prevent unauthorized physical access, damage and interference to the organization\'s '
                'information and other associated assets in offices, rooms and facilities.'
            ),
            'observation': (
                'The server room at Yeti Airlines\' head office houses critical infrastructure including booking '
                'system servers, email servers, and the IVR/telephony system. The following observations were made: '
                '(1) The server room has a raised floor but lacks proper cable management, creating trip hazards and '
                'making maintenance difficult. (2) Environmental controls are present (air conditioning), but no '
                'environmental monitoring sensors (temperature, humidity, water leak detection) with alerting '
                'capabilities were found. (3) Fire suppression in the server room relies on standard fire '
                'extinguishers rather than a gas-based suppression system appropriate for electronic equipment. '
                '(4) UPS systems are in place but their capacity and maintenance records were not available for review. '
                '(5) The DRC facility does not have redundant power supply arrangements (no diesel generator backup '
                'was observed).'
            ),
            'risk_rating': 'MEDIUM',
            'recommendation': (
                'Yeti Airlines should implement environmental monitoring systems with real-time alerting for '
                'temperature, humidity, and water leak detection in both the DC and DRC. A gas-based fire suppression '
                'system (such as FM-200 or Novec 1230) should be installed in the server room. Proper cable management '
                'should be implemented following structured cabling standards. UPS capacity should be reviewed and '
                'tested regularly, with maintenance records properly documented. The DRC should be equipped with a '
                'diesel generator as a secondary power backup. An Auto Transfer Switch (ATS) should be configured '
                'for seamless power failover.'
            ),
            'management_response': 'The IT department has initiated a project to revamp the server room. Environmental monitoring sensors and gas-based fire suppression systems will be installed by the end of the year. The DRC power backup requirements are under review.',
        },
        {
            'section': 'A.4. Technological Controls',
            'subsection': 'A.4.1. Access Control (A.8.2)',
            'control': (
                'Access to information and other associated assets should be restricted in accordance with the '
                'established topic-specific policy on access control.'
            ),
            'purpose': (
                'To ensure authorized access and to prevent unauthorized access to information and other '
                'associated assets.'
            ),
            'observation': (
                'Access control at Yeti Airlines was found to have significant gaps: (1) A formal access control '
                'policy has not been documented. (2) User account management lacks a structured provisioning and '
                'de-provisioning process; terminated employees\' accounts were found to remain active in some '
                'systems for extended periods. (3) Shared accounts were observed in the flight booking system at '
                'multiple counter locations. (4) Password policy enforcement is weak\u2014minimum length requirements '
                'are set to 6 characters with no complexity rules, and password expiry is not enforced. '
                '(5) Privileged Access Management (PAM) is not implemented; system administrators use shared '
                'root/admin credentials. (6) Multi-factor authentication (MFA) is not implemented for any '
                'critical system, including the booking portal and email system. (7) Remote access to internal '
                'systems by IT staff is done via basic VPN without additional authentication controls.'
            ),
            'risk_rating': 'HIGH',
            'recommendation': (
                'Yeti Airlines should develop and implement a formal access control policy. A structured user '
                'lifecycle management process (provisioning, review, de-provisioning) should be established with '
                'HR integration for timely account deactivation upon employee separation. Shared accounts should '
                'be eliminated and replaced with individual user accounts with proper audit trails. Password '
                'policy should enforce minimum 12-character passwords with complexity requirements and regular '
                'rotation. A Privileged Access Management (PAM) solution should be implemented. Multi-factor '
                'authentication (MFA) should be mandated for all critical systems, including email, booking '
                'systems, and remote VPN access. Periodic user access reviews should be conducted at least '
                'quarterly.'
            ),
            'management_response': 'An access control policy is currently being drafted. We will enforce stronger password policies and eliminate shared accounts immediately. The implementation of PAM and MFA for critical systems is planned for the next fiscal year.',
        },
        {
            'subsection': 'A.4.2. Protection Against Malware (A.8.7)',
            'control': (
                'Protection against malware should be implemented and supported by appropriate user awareness.'
            ),
            'purpose': (
                'To ensure information and other associated assets are protected against malware.'
            ),
            'observation': (
                'Yeti Airlines has deployed antivirus software on most workstations at the head office; however, '
                'the following gaps were observed: (1) The antivirus solution is not centrally managed, making it '
                'difficult to ensure consistent deployment and updates across all endpoints. (2) Several '
                'workstations at branch offices (particularly Nepalgunj and Bhairahawa) were found running '
                'outdated virus definitions (more than 30 days old). (3) No endpoint detection and response (EDR) '
                'solution is deployed on any system. (4) USB port restrictions are not enforced, and removable '
                'media usage is common among staff for data transfer. (5) Email gateway filtering is basic, with '
                'no advanced threat protection or sandboxing capability. (6) No regular malware scanning schedule '
                'has been established for servers.'
            ),
            'risk_rating': 'MEDIUM',
            'recommendation': (
                'Yeti Airlines should implement a centrally managed endpoint protection platform that provides '
                'real-time visibility and control across all endpoints, including branch offices. The solution '
                'should include EDR capabilities for advanced threat detection. USB port restrictions should be '
                'enforced through group policy or endpoint management tools. Email security should be enhanced '
                'with advanced threat protection including sandboxing and URL filtering. Regular full-system '
                'malware scans should be scheduled on all servers and workstations. An approved media policy '
                'should be developed and communicated to all staff.'
            ),
            'management_response': 'We are evaluating centralized endpoint protection platforms with EDR capabilities. USB restrictions will be enforced via group policy next month. An updated removable media policy will also be published.',
        },
        {
            'subsection': 'A.4.3. Management of Technical Vulnerabilities (A.8.8)',
            'control': (
                'Information about technical vulnerabilities of information systems in use should be obtained, '
                'the organization\'s exposure to such vulnerabilities should be evaluated and appropriate '
                'measures should be taken.'
            ),
            'purpose': (
                'To prevent exploitation of technical vulnerabilities.'
            ),
            'observation': (
                'Yeti Airlines does not have a formal vulnerability management program in place. The following '
                'gaps were identified: (1) No periodic vulnerability assessments or penetration tests are '
                'conducted on IT infrastructure or web applications (including the online booking portal and '
                'mobile application). (2) Patch management is ad-hoc, with no defined process or timeline for '
                'applying security patches. Several servers were found running outdated operating system versions '
                'with known vulnerabilities. (3) The airline\'s customer-facing web application has not undergone '
                'a security assessment since its deployment. (4) No software inventory or asset register is '
                'maintained to track software versions and patch status. (5) There is no subscription to '
                'vulnerability intelligence feeds or advisory services relevant to the technologies in use.'
            ),
            'risk_rating': 'HIGH',
            'recommendation': (
                'Yeti Airlines should establish a formal vulnerability management program that includes regular '
                'vulnerability assessments (at least quarterly) and annual penetration testing for all critical '
                'systems, including the online booking portal, mobile application, and internal IT infrastructure. '
                'A structured patch management process should be defined with clear timelines for applying critical '
                'patches (within 72 hours), high-priority patches (within 2 weeks), and routine patches (within '
                '30 days). A comprehensive IT asset inventory should be maintained. The airline should subscribe '
                'to relevant vulnerability advisory services (e.g., CERT-NP, vendor advisories) for timely '
                'awareness of emerging threats.'
            ),
            'management_response': 'Management will engage a third-party security firm to conduct annual penetration testing starting this year. A formal patch management process and a comprehensive IT asset inventory system are currently being developed.',
        },
        {
            'subsection': 'A.4.4. Information Backup (A.8.13)',
            'control': (
                'Backup copies of information, software and systems should be maintained and regularly tested '
                'in accordance with the agreed topic-specific policy on backup.'
            ),
            'purpose': (
                'To enable recovery of information and other associated assets following data loss or disruption.'
            ),
            'observation': (
                'Yeti Airlines performs daily backups of critical databases (booking system, financial system); '
                'however, the following concerns were identified: (1) No formal backup policy exists that defines '
                'backup scope, frequency, retention periods, and recovery procedures. (2) Backups are stored '
                'on-site at the DC, but off-site backup copies at the DRC are only replicated weekly, creating '
                'a potential data loss window of up to 7 days. (3) Backup restoration tests are not performed '
                'regularly\u2014the last documented restoration test was conducted over 12 months ago. (4) Recovery '
                'Time Objective (RTO) and Recovery Point Objective (RPO) have not been formally defined for '
                'critical systems. (5) Backup encryption is not implemented, and backup media is not stored in a '
                'fireproof safe. (6) There is no Business Continuity Plan (BCP) that integrates backup and '
                'recovery procedures with broader continuity objectives.'
            ),
            'risk_rating': 'HIGH',
            'recommendation': (
                'Yeti Airlines should develop a formal backup policy defining RTO, RPO, backup frequency, '
                'retention periods, and restoration procedures for all critical systems. Backup replication to '
                'the DRC should be performed daily at minimum, with real-time replication considered for the '
                'booking system. Quarterly backup restoration tests should be conducted and documented. All '
                'backup data should be encrypted both in transit and at rest. A comprehensive Business Continuity '
                'Plan (BCP) should be developed, tested annually, and integrated with the backup strategy. '
                'Backup media stored off-site should be secured in fireproof, climate-controlled storage.'
            ),
            'management_response': 'The IT department will draft a formal backup policy defining RTO and RPO within three months. We will also upgrade our backup infrastructure to support daily automated replication to the DRC and implement backup encryption.',
        },
        {
            'subsection': 'A.4.5. Logging (A.8.15)',
            'control': (
                'Logs that record activities, exceptions, faults and other relevant events should be produced, '
                'stored, protected and analysed.'
            ),
            'purpose': (
                'To record events, generate evidence, ensure the integrity of log information, prevent '
                'unauthorized access, identify information security events that may lead to an information '
                'security incident, and to support investigations.'
            ),
            'observation': (
                'Logging and monitoring at Yeti Airlines is rudimentary and fragmented: (1) System logs are '
                'generated on individual servers and network devices, but there is no centralized log management '
                'or Security Information and Event Management (SIEM) solution in place. (2) Log retention is '
                'inconsistent\u2014some systems retain logs for 30 days while others overwrite logs after 7 days. '
                '(3) No log review process is established; logs are only examined reactively in the event of a '
                'reported issue. (4) Audit logging for the booking system does not capture all critical events '
                '(e.g., booking modifications, cancellations, refund processing). (5) Administrator activities '
                'are not separately logged or monitored. (6) Logs are not protected against tampering\u2014system '
                'administrators with root access can modify or delete log files. (7) No alerting mechanism '
                'exists for security-relevant events such as multiple failed login attempts, privilege escalation, '
                'or after-hours access.'
            ),
            'risk_rating': 'HIGH',
            'recommendation': (
                'Yeti Airlines should implement a centralized log management solution or SIEM platform to '
                'aggregate, correlate, and analyze logs from all critical systems and network devices. A minimum '
                'log retention period of 1 year should be established for security-relevant logs. Automated '
                'alerting should be configured for critical security events including failed login attempts, '
                'privilege escalation, unauthorized access attempts, and configuration changes. Logs should be '
                'write-protected and stored in a tamper-evident manner. Administrator activities should be logged '
                'separately with enhanced monitoring. Regular log reviews should be conducted weekly at minimum. '
                'The logging framework should be aligned with CAAN regulatory requirements and the Electronic '
                'Transactions Act (2008) evidence preservation requirements.'
            ),
            'management_response': 'The acquisition of a centralized log management and SIEM solution is included in the IT budget for the next financial year. In the interim, log retention policies will be standardized across all critical servers.',
        },
    ]

    # ==================== RENDER CONTROLS ====================
    current_section = None
    for ctrl in controls:
        if 'section' in ctrl and ctrl['section'] != current_section:
            current_section = ctrl['section']
            h = doc.add_heading(current_section, level=2)
            style_heading(h)

        h = doc.add_heading(ctrl['subsection'], level=3)
        style_heading(h)

        # Control
        p = doc.add_paragraph()
        run = p.add_run('Control')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = BLACK
        doc.add_paragraph(ctrl['control'])

        # Purpose
        p = doc.add_paragraph()
        run = p.add_run('Purpose')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = BLACK
        doc.add_paragraph(ctrl['purpose'])

        # Observations
        p = doc.add_paragraph()
        run = p.add_run('Observations')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = BLACK
        doc.add_paragraph(ctrl['observation'])

        # Risk Rating
        p = doc.add_paragraph()
        run = p.add_run('Risk Rating: ')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = BLACK
        run = p.add_run(ctrl['risk_rating'])
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = BLACK

        # Recommendation
        p = doc.add_paragraph()
        run = p.add_run('Recommendation')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = BLACK
        doc.add_paragraph(ctrl['recommendation'])

        # Management Response
        p = doc.add_paragraph()
        run = p.add_run('Management Response')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = BLACK
        doc.add_paragraph(ctrl['management_response'])

    doc.add_page_break()

    # ==================== SECTION B: SUMMARY OF FINDINGS ====================
    h = doc.add_heading('B. Summary of Findings', level=1)
    style_heading(h)

    doc.add_paragraph(
        'The following table summarizes the findings from the IS Audit of Yeti Airlines Pvt. Ltd. '
        'based on the 10 selected ISO 27001:2022 controls:'
    )

    summary_table = doc.add_table(rows=11, cols=4)
    summary_table.style = 'Table Grid'

    headers = ['S.N.', 'Control Area', 'ISO Reference', 'Risk Rating']
    for i, header in enumerate(headers):
        cell = summary_table.cell(0, i)
        cell.text = ''
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = BLACK

    summary_data = [
        ('1', 'Policies for Information Security', 'A.5.1', 'HIGH'),
        ('2', 'Information Security Roles and Responsibilities', 'A.5.2', 'HIGH'),
        ('3', 'Information Security Awareness, Education and Training', 'A.6.3', 'HIGH'),
        ('4', 'Physical Security Perimeters', 'A.7.1', 'HIGH'),
        ('5', 'Securing Offices, Rooms and Facilities', 'A.7.3', 'MEDIUM'),
        ('6', 'Access Control', 'A.8.2', 'HIGH'),
        ('7', 'Protection Against Malware', 'A.8.7', 'MEDIUM'),
        ('8', 'Management of Technical Vulnerabilities', 'A.8.8', 'HIGH'),
        ('9', 'Information Backup', 'A.8.13', 'HIGH'),
        ('10', 'Logging', 'A.8.15', 'HIGH'),
    ]

    for row_idx, (sn, area, ref, rating) in enumerate(summary_data, 1):
        for j, text in enumerate([sn, area, ref, rating]):
            cell = summary_table.cell(row_idx, j)
            cell.text = ''
            run = cell.paragraphs[0].add_run(text)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = BLACK

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('Overall Assessment: ')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = BLACK
    run = p.add_run(
        'Out of 10 controls evaluated, 8 controls were rated as HIGH risk and 2 controls were rated as '
        'MEDIUM risk. This indicates that Yeti Airlines\' information security posture requires significant '
        'improvement across organizational, physical, and technological domains. Immediate attention is '
        'required to address the identified gaps, particularly in access control, vulnerability management, '
        'backup and recovery, and logging and monitoring areas.'
    )
    run.font.name = 'Times New Roman'
    run.font.color.rgb = BLACK

    doc.add_page_break()

    # ==================== SECTION C: RECOMMENDATIONS ====================
    h = doc.add_heading('C. Recommendations', level=1)
    style_heading(h)

    doc.add_paragraph(
        'This section provides a prioritized list of recommendations based on the observations and '
        'evidence gathered during the course of the Information System Audit at Yeti Airlines Pvt. Ltd. '
        'The recommendations are categorized by priority level.'
    )

    # Critical Priority
    p = doc.add_paragraph()
    run = p.add_run('Critical Priority (Immediate Action Required)')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = BLACK

    critical_recs = [
        'Implement Multi-Factor Authentication (MFA) for all critical systems including the online booking '
        'portal, email system, flight operations systems, and remote VPN access.',
        'Eliminate shared user accounts across all systems and implement individual user accounts with '
        'proper audit trails.',
        'Conduct an immediate vulnerability assessment and penetration test of the customer-facing online '
        'booking portal and mobile application to identify and remediate critical vulnerabilities.',
        'Implement a centralized log management / SIEM solution to enable real-time monitoring and alerting '
        'of security events.',
        'Establish formal RTO and RPO definitions for all critical systems and increase DRC backup '
        'replication frequency from weekly to daily.',
    ]
    for rec in critical_recs:
        p = doc.add_paragraph(rec, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    # High Priority
    p = doc.add_paragraph()
    run = p.add_run('High Priority (Within 3 Months)')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = BLACK

    high_recs = [
        'Review and update the Information Security Policy to align with ISO 27001:2022 and the current '
        'organizational context. Develop topic-specific policies for acceptable use, data classification, '
        'access control, and remote access.',
        'Appoint a dedicated Chief Information Security Officer (CISO) or Information Security Manager '
        'with clear authority and reporting lines.',
        'Implement electronic access control systems (biometric or smart card) for the Data Center and '
        'Disaster Recovery Center.',
        'Deploy a centrally managed endpoint protection platform with EDR capabilities across all '
        'workstations and servers, including branch offices.',
        'Establish a formal vulnerability management and patch management program with defined timelines '
        'for applying security patches.',
        'Develop and implement a comprehensive Business Continuity Plan (BCP) and conduct a tabletop '
        'exercise to test its effectiveness.',
    ]
    for rec in high_recs:
        p = doc.add_paragraph(rec, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    # Medium Priority
    p = doc.add_paragraph()
    run = p.add_run('Medium Priority (Within 6 Months)')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = BLACK

    medium_recs = [
        'Provide cybersecurity awareness training to all employees of Yeti Airlines at planned intervals '
        '(at least annually), covering phishing, social engineering, password hygiene, and incident reporting.',
        'Install environmental monitoring systems (temperature, humidity, water leak detection) with '
        'real-time alerting in the DC and DRC.',
        'Install a gas-based fire suppression system in the server room.',
        'Implement Privileged Access Management (PAM) solution for managing administrative and root-level access.',
        'Establish security focal points at all major stations (Pokhara, Nepalgunj, Bhairahawa, Biratnagar, '
        'Bhadrapur) with defined escalation procedures.',
        'Perform a comprehensive Risk Assessment for all IT assets and projects at planned intervals.',
        'Ensure compliance with Nepal\'s Individual Privacy Act (2018) and Electronic Transactions Act (2008) '
        'in all IT systems processing passenger personal data.',
    ]
    for rec in medium_recs:
        p = doc.add_paragraph(rec, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # ==================== REGULATORY COMPLIANCE ====================
    h = doc.add_heading('D. Regulatory Compliance Context', level=1)
    style_heading(h)

    doc.add_paragraph(
        'Yeti Airlines Pvt. Ltd., as a domestic airline operating under the jurisdiction of the Civil Aviation '
        'Authority of Nepal (CAAN), is subject to the following regulatory requirements that have implications '
        'for information security:'
    )

    reg_items = [
        'Civil Aviation Act, 2053 (1996) \u2013 Governs overall aviation operations and safety requirements.',
        'Civil Aviation Security Regulation, 2073 (2016) \u2013 Mandates security measures for civil aviation, '
        'including information and communication security aspects.',
        'Electronic Transactions Act (ETA), 2063 (2008) \u2013 The foundational cyber-law that criminalizes '
        'unauthorized access, data theft, and electronic fraud. Relevant to the airline\'s online booking '
        'and payment processing systems.',
        'Individual Privacy Act, 2075 (2018) \u2013 Governs the collection, processing, storage, and use of '
        'personal data, directly applicable to passenger data handling.',
        'Individual Privacy Regulation, 2077 (2020) \u2013 Implementation procedures for the Privacy Act.',
        'ICAO Standards and Recommended Practices (SARPs) \u2013 As a signatory to the Chicago Convention, '
        'Nepal aligns aviation security with ICAO Annex 17 standards.',
        'IATA Standard Safety Assessment (ISSA) \u2013 Yeti Airlines has obtained ISSA certification, which '
        'includes requirements for safety management systems.',
        'NTA Cyber Security Bylaw (2020) \u2013 While primarily for telecom operators, its standards serve as '
        'reference for cybersecurity best practices in Nepal.',
    ]
    for item in reg_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()
    doc.add_paragraph()

    # ==================== END OF REPORT ====================
    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_para.paragraph_format.space_before = Pt(36)
    run = end_para.add_run('END OF REPORT')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = BLACK

    # Format all tables: remove line spacing, add padding top and bottom
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(3)
                    paragraph.paragraph_format.space_after = Pt(3)

    return doc


if __name__ == '__main__':
    doc = create_document()
    output_path = '/Users/bidur/ncit/isa/ME_252952_Bidur_Sapkota_IS_Audit_Report_Yeti_Airlines.docx'
    doc.save(output_path)
    print(f'Report generated successfully: {output_path}')
